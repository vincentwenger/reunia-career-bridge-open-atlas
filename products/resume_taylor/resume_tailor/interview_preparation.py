from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable, Literal

from docx import Document
from pydantic import Field, field_validator

from .grounding import validate_candidate_claim
from .models import CandidateProfile, StrictModel
from .resume_findings import ResumeFindingsSnapshot
from .web_state import WorkflowState, normalize_job_description


class InterviewPreparationPoint(StrictModel):
    title: str
    detail: str
    evidence_ids: list[str] = Field(default_factory=list)

    @field_validator("title", "detail")
    @classmethod
    def normalize_text(cls, value: str) -> str:
        return " ".join(value.split())


class InterviewQuestionPlan(StrictModel):
    question: str
    why_likely: str
    answer_focus: str
    evidence_ids: list[str] = Field(default_factory=list)

    @field_validator("question", "why_likely", "answer_focus")
    @classmethod
    def normalize_text(cls, value: str) -> str:
        return " ".join(value.split())


GapEvidenceStatus = Literal["partial_verified_evidence", "no_verified_evidence"]


class InterviewExperienceGap(StrictModel):
    requirement: str
    evidence_status: GapEvidenceStatus
    explanation: str
    preparation_action: str

    @field_validator("requirement", "explanation", "preparation_action")
    @classmethod
    def normalize_text(cls, value: str) -> str:
        return " ".join(value.split())


class PersonalIntroductionOutline(StrictModel):
    opening: str
    current_value: str
    relevant_background: str
    role_connection: str
    closing: str
    evidence_ids: list[str] = Field(default_factory=list)

    @field_validator(
        "opening", "current_value", "relevant_background", "role_connection", "closing"
    )
    @classmethod
    def normalize_text(cls, value: str) -> str:
        return " ".join(value.split())


class InterviewPreparationWorkspace(StrictModel):
    role_summary: str
    company_summary: str
    expected_responsibilities: list[str] = Field(default_factory=list)
    likely_technical_questions: list[InterviewQuestionPlan] = Field(default_factory=list)
    likely_behavioral_questions: list[InterviewQuestionPlan] = Field(default_factory=list)
    resume_challenge_areas: list[InterviewPreparationPoint] = Field(default_factory=list)
    candidate_strengths: list[InterviewPreparationPoint] = Field(default_factory=list)
    potential_experience_gaps: list[InterviewExperienceGap] = Field(default_factory=list)
    questions_to_ask: list[str] = Field(default_factory=list)
    personal_introduction: PersonalIntroductionOutline

    @field_validator("role_summary", "company_summary")
    @classmethod
    def normalize_summary(cls, value: str) -> str:
        return " ".join(value.split())

    @field_validator("expected_responsibilities", "questions_to_ask")
    @classmethod
    def normalize_list(cls, values: list[str]) -> list[str]:
        seen: set[str] = set()
        normalized: list[str] = []
        for raw in values:
            value = " ".join(str(raw).split())
            key = value.casefold()
            if value and key not in seen:
                normalized.append(value)
                seen.add(key)
        return normalized


@dataclass(frozen=True)
class VerifiedEvidenceItem:
    id: str
    text: str
    source: str


@dataclass(frozen=True)
class VerifiedEvidenceBundle:
    items: tuple[VerifiedEvidenceItem, ...]
    source_label: str
    fingerprint: str

    @property
    def count(self) -> int:
        return len(self.items)

    @property
    def ids(self) -> frozenset[str]:
        return frozenset(item.id for item in self.items)

    @property
    def submitted_resume_ids(self) -> frozenset[str]:
        return frozenset(
            item.id
            for item in self.items
            if item.source == "submitted evidence-reviewed resume"
        )

    def prompt_text(self) -> str:
        return "\n".join(
            f"[{item.id}] ({item.source}) {item.text}" for item in self.items
        )


INTERVIEW_PREPARATION_SYSTEM = """You create an evidence-protected interview preparation workspace for one specific job application.

Grounding rules:
1. Use the target job description to determine the role, responsibilities, likely questions, and gaps.
2. Use only the supplied VERIFIED CANDIDATE EVIDENCE for statements about the candidate.
3. Every candidate strength, resume challenge, answer recommendation, and personal-introduction claim must cite one or more supplied evidence IDs.
4. Never invent experience, seniority, years, metrics, tools, credentials, employers, industries, responsibilities, or accomplishments.
5. Potential gaps must be described as an absence or partial match in the supplied evidence, not as a definitive claim about everything the candidate has ever done.
6. The company summary must use only the company name and facts present in the job description. Do not add outside company research. When the posting provides little company information, say so plainly.
7. Do not ask for or discuss immigration, visa, nationality, protected characteristics, compensation history, or other unnecessary sensitive information.
8. Keep the preparation practical and specific. Include approximately 5-7 technical questions and 5-7 behavioral questions when the job description supports them.
9. Questions to ask the interviewer should help the candidate evaluate responsibilities, success measures, team practices, and role expectations; avoid questions already answered clearly in the posting.
10. The personal introduction is an outline, not a fabricated script. It should be concise enough for a 60-90 second response.
11. Treat the supplied STRUCTURED RESUME FINDINGS as authoritative workflow findings, but never treat a finding by itself as proof of candidate experience. Candidate claims still require verified evidence IDs.
12. Do not revive excluded, unsupported, or unanswered claims. Use them only to prepare honest explanations, challenge questions, or gap-bridging actions.

Return only the requested structured result."""


def build_interview_preparation_prompt(
    *,
    company: str,
    role: str,
    job_description: str,
    evidence: VerifiedEvidenceBundle,
    resume_findings: ResumeFindingsSnapshot,
    interview_audience: str = "",
    career_profile_context: dict[str, str] | None = None,
) -> str:
    profile_json = json.dumps(
        career_profile_context or {}, ensure_ascii=False, indent=2, sort_keys=True
    )
    return f"""Build the interview preparation workspace below.

APPLICATION
Company: {company.strip() or 'Not specified'}
Target role: {role.strip() or 'Not specified'}
Interview audience: {interview_audience.strip() or 'Not specified'}

TARGET JOB DESCRIPTION
---
{normalize_job_description(job_description)}
---

REUSABLE CAREER PROFILE — CONTEXT ONLY
Use this to choose relevant topics, explain international background, respect career goals and constraints, and tailor questions to the candidate's direction. It is not verified evidence. Never turn an accomplishment, skill, credential, title, year count, or work-authorization statement from this section into a candidate claim unless a VERIFIED CANDIDATE EVIDENCE ID supports it.
---
{profile_json}
---

VERIFIED CANDIDATE EVIDENCE
Evidence source: {evidence.source_label}
Every candidate claim must be traceable to one or more of these exact IDs.
---
{evidence.prompt_text()}
---

STRUCTURED RESUME FINDINGS
These findings were saved by the resume workflow for this application. Use them directly to prioritize interview preparation.
---
{resume_findings.prompt_text()}
---

HOW TO USE THE FINDINGS
- Convert unsupported or partial requirements into explicit potential gaps, likely probing questions, and honest bridge plans.
- Convert evidence-review warnings and resume-report weaknesses into targeted preparation or resume challenge areas when relevant to an interview.
- Use Career Translation Assessment findings to prepare explanations of unfamiliar international titles, credentials, terminology, and transferable experience.
- Use alignment changes to emphasize what improved while still preparing for the remaining weak areas.
- Use excluded or questioned claims only as warnings about what must not be overstated or as prompts for clarification. Never present them as verified experience.

SECTION REQUIREMENTS
- Role summary: explain the role's purpose in 2-4 sentences.
- Company summary: summarize only what the posting says about the organization, business, customers, mission, or team.
- Expected responsibilities: provide the central responsibilities from the posting.
- Likely technical questions: include why each question is likely, an evidence-safe answer focus, and relevant evidence IDs. When no evidence supports an answer, keep evidence_ids empty and recommend an honest bridge rather than inventing experience.
- Likely behavioral questions: connect each answer focus to verified evidence IDs.
- Resume challenge areas: identify resume statements, transitions, unfamiliar titles, scope, metrics, or technical claims an interviewer may probe. When evidence items labeled "submitted evidence-reviewed resume" are available, cite only those IDs in this section.
- Candidate strengths: identify the strongest verified matches to emphasize and cite evidence IDs.
- Potential experience gaps: compare the posting with supplied evidence and label each as partial_verified_evidence or no_verified_evidence.
- Questions to ask: provide thoughtful, role-specific interviewer questions.
- Personal introduction: create a five-part outline based only on verified evidence and cite all evidence IDs used.
"""


def _clean_id_fragment(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip()).strip("-")
    return cleaned[:64] or "item"


def _profile_evidence(profile: CandidateProfile) -> list[VerifiedEvidenceItem]:
    items: list[VerifiedEvidenceItem] = []
    if profile.current_summary.strip():
        items.append(
            VerifiedEvidenceItem("summary", profile.current_summary.strip(), "verified profile summary")
        )

    skill_groups = (
        ("hard", profile.skills.hard_skills),
        ("soft", profile.skills.soft_skills),
        ("tool", profile.skills.tools_software),
        ("industry", profile.skills.industry_knowledge),
        ("language", profile.skills.languages),
    )
    for group, skills in skill_groups:
        for index, skill in enumerate(skills, start=1):
            if skill.strip():
                items.append(
                    VerifiedEvidenceItem(
                        f"skill-{group}-{index}", skill.strip(), f"verified {group} skill"
                    )
                )

    for experience in profile.experiences:
        experience_id = _clean_id_fragment(experience.id)
        heading = " · ".join(
            part for part in (experience.title, experience.employer, experience.dates) if part.strip()
        )
        if heading:
            items.append(
                VerifiedEvidenceItem(
                    f"experience-{experience_id}", heading, "verified employment record"
                )
            )
        for bullet in experience.bullets:
            if bullet.text.strip():
                items.append(
                    VerifiedEvidenceItem(
                        _clean_id_fragment(bullet.id),
                        bullet.text.strip(),
                        f"verified experience: {experience.title} at {experience.employer}",
                    )
                )

    for index, education in enumerate(profile.education, start=1):
        text = " · ".join(
            part
            for part in (
                education.credential,
                education.institution,
                education.location,
                education.date,
                education.detail,
            )
            if part.strip()
        )
        if text:
            items.append(
                VerifiedEvidenceItem(f"education-{index}", text, "verified education")
            )

    for evidence_item in profile.supplemental_evidence:
        if evidence_item.statement.strip():
            items.append(
                VerifiedEvidenceItem(
                    _clean_id_fragment(evidence_item.id),
                    evidence_item.statement.strip(),
                    "candidate-confirmed supplemental evidence",
                )
            )
    return items


def _unique_evidence_items(
    items: Iterable[VerifiedEvidenceItem],
) -> list[VerifiedEvidenceItem]:
    counts: dict[str, int] = {}
    unique: list[VerifiedEvidenceItem] = []
    for item in items:
        base_id = _clean_id_fragment(item.id)
        counts[base_id] = counts.get(base_id, 0) + 1
        unique_id = base_id if counts[base_id] == 1 else f"{base_id}-{counts[base_id]}"
        unique.append(VerifiedEvidenceItem(unique_id, item.text, item.source))
    return unique


def _docx_evidence(resume_bytes: bytes) -> list[VerifiedEvidenceItem]:
    try:
        document = Document(BytesIO(resume_bytes))
    except Exception:
        return []

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            text = " | ".join(
                " ".join(cell.text.split()) for cell in row.cells if cell.text.strip()
            )
            if text:
                lines.append(text)

    seen: set[str] = set()
    items: list[VerifiedEvidenceItem] = []
    for text in lines:
        key = text.casefold()
        if key in seen:
            continue
        seen.add(key)
        items.append(
            VerifiedEvidenceItem(
                f"submitted-resume-{len(items) + 1:03d}",
                text,
                "submitted evidence-reviewed resume",
            )
        )
    return items


def build_verified_evidence_bundle(
    workflow_state: WorkflowState,
    *,
    submitted_resume_bytes: bytes | None,
) -> VerifiedEvidenceBundle:
    profile = workflow_state.final_report_profile or workflow_state.confirmed_profile
    profile_is_verified = profile is not None and (
        workflow_state.confirmation_complete or workflow_state.final_report_profile is not None
    )
    profile_items = _profile_evidence(profile) if profile_is_verified and profile else []
    resume_items = _docx_evidence(submitted_resume_bytes) if submitted_resume_bytes else []
    items = _unique_evidence_items([*profile_items, *resume_items])

    if profile_items and resume_items:
        source_label = (
            "Confirmed Candidate Profile plus the submitted evidence-reviewed resume "
            "attached to this application"
        )
    elif profile_items:
        source_label = "Confirmed Candidate Profile and evidence accepted in the resume workflow"
    elif resume_items:
        source_label = "Submitted evidence-reviewed resume attached to this application"
    else:
        source_label = "No verified candidate evidence available"

    payload = {
        "source": source_label,
        "items": [{"id": item.id, "text": item.text, "source": item.source} for item in items],
    }
    fingerprint = hashlib.sha256(
        json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()
    return VerifiedEvidenceBundle(tuple(items), source_label, fingerprint)


def job_description_fingerprint(
    job_description: str,
    *,
    company: str = "",
    role: str = "",
    interview_audience: str = "",
    career_profile_fingerprint: str = "",
) -> str:
    payload = {
        "company": " ".join(company.split()),
        "role": " ".join(role.split()),
        "interview_audience": " ".join(interview_audience.split()),
        "career_profile_fingerprint": str(career_profile_fingerprint or "").strip(),
        "job_description": normalize_job_description(job_description),
    }
    return hashlib.sha256(
        json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()


def restrict_workspace_to_evidence(
    workspace: InterviewPreparationWorkspace,
    allowed_ids: Iterable[str],
    *,
    submitted_resume_ids: Iterable[str] = (),
    evidence_by_id: dict[str, str] | None = None,
) -> InterviewPreparationWorkspace:
    allowed = frozenset(allowed_ids)
    submitted = frozenset(submitted_resume_ids)
    evidence_lookup = {
        str(evidence_id): str(text)
        for evidence_id, text in (evidence_by_id or {}).items()
        if str(evidence_id).strip() and str(text).strip()
    }

    def clean(ids: list[str], accepted: frozenset[str] = allowed) -> list[str]:
        seen: set[str] = set()
        cleaned: list[str] = []
        for value in ids:
            candidate = value.strip().strip("[](){}<>").strip()
            if candidate in accepted and candidate not in seen:
                cleaned.append(candidate)
                seen.add(candidate)
        return cleaned

    def is_grounded(text: str, ids: list[str], *, require_overlap: bool = True) -> bool:
        if not evidence_lookup:
            return True
        cited_text = [evidence_lookup[evidence_id] for evidence_id in ids if evidence_id in evidence_lookup]
        if not cited_text:
            return False
        return not validate_candidate_claim(
            text,
            cited_text,
            require_overlap=require_overlap,
        )

    for question in workspace.likely_technical_questions:
        question.evidence_ids = clean(question.evidence_ids)
        if question.evidence_ids and not is_grounded(
            question.answer_focus,
            question.evidence_ids,
            require_overlap=True,
        ):
            question.evidence_ids = []
            question.answer_focus = (
                "No verified evidence supports a specific candidate claim; prepare an honest "
                "bridge and do not imply prior experience."
            )

    grounded_behavioral: list[InterviewQuestionPlan] = []
    for question in workspace.likely_behavioral_questions:
        question.evidence_ids = clean(question.evidence_ids)
        if question.evidence_ids and is_grounded(
            question.answer_focus,
            question.evidence_ids,
            require_overlap=True,
        ):
            grounded_behavioral.append(question)
    workspace.likely_behavioral_questions = grounded_behavioral

    grounded_challenges: list[InterviewPreparationPoint] = []
    challenge_ids = submitted or allowed
    for point in workspace.resume_challenge_areas:
        point.evidence_ids = clean(point.evidence_ids, challenge_ids)
        if (
            point.evidence_ids
            and is_grounded(point.title, point.evidence_ids, require_overlap=False)
            and is_grounded(point.detail, point.evidence_ids, require_overlap=True)
        ):
            grounded_challenges.append(point)
    workspace.resume_challenge_areas = grounded_challenges

    grounded_strengths: list[InterviewPreparationPoint] = []
    for point in workspace.candidate_strengths:
        point.evidence_ids = clean(point.evidence_ids)
        if (
            point.evidence_ids
            and is_grounded(point.title, point.evidence_ids, require_overlap=False)
            and is_grounded(point.detail, point.evidence_ids, require_overlap=True)
        ):
            grounded_strengths.append(point)
    workspace.candidate_strengths = grounded_strengths

    workspace.personal_introduction.evidence_ids = clean(
        workspace.personal_introduction.evidence_ids
    )
    introduction_text = " ".join(
        (
            workspace.personal_introduction.opening,
            workspace.personal_introduction.current_value,
            workspace.personal_introduction.relevant_background,
            workspace.personal_introduction.role_connection,
            workspace.personal_introduction.closing,
        )
    )
    if workspace.personal_introduction.evidence_ids and not is_grounded(
        introduction_text,
        workspace.personal_introduction.evidence_ids,
        require_overlap=True,
    ):
        workspace.personal_introduction.evidence_ids = []
    if not workspace.personal_introduction.evidence_ids:
        raise ValueError(
            "The generated personal introduction did not contain traceable verified evidence."
        )
    return workspace
