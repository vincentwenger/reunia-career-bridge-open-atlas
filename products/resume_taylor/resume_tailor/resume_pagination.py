from __future__ import annotations

from dataclasses import dataclass
from math import ceil
from typing import Iterable

from docx.shared import Pt

from .docx_styles import (
    STYLE_BULLET,
    STYLE_CONTACT,
    STYLE_EDUCATION,
    STYLE_EDUCATION_DETAIL,
    STYLE_EMPLOYER_LINE,
    STYLE_NAME,
    STYLE_ROLE,
    STYLE_SECTION_HEADING,
    STYLE_SKILL_LINE,
    STYLE_SUMMARY,
    STYLE_TARGET_TITLE,
)

# Word and LibreOffice use font metrics, grid rounding, and printer-layout
# tolerances that python-docx cannot calculate. Reserving one typographic line
# keeps the pure-Python estimate conservative enough to catch the common case
# where a single credential or bullet is pushed to an otherwise empty page.
PAGE_LAYOUT_SAFETY_RESERVE_POINTS = 14.0
ORPHAN_LAST_PAGE_MIN_LINES = 3
ORPHAN_LAST_PAGE_MIN_FILL_RATIO = 0.20


@dataclass(frozen=True)
class ParagraphLayoutEstimate:
    index: int
    height_points: float
    substantive_lines: int
    keep_with_next: bool
    page_break_before: bool
    style_name: str


@dataclass(frozen=True)
class PaginationEstimate:
    page_count: int
    used_points_by_page: tuple[float, ...]
    lines_by_page: tuple[int, ...]
    usable_height_points: float

    @property
    def last_page_fill_ratio(self) -> float:
        if not self.used_points_by_page or self.usable_height_points <= 0:
            return 0.0
        return min(1.0, self.used_points_by_page[-1] / self.usable_height_points)

    @property
    def last_page_substantive_lines(self) -> int:
        return self.lines_by_page[-1] if self.lines_by_page else 0

    @property
    def has_orphan_final_page(self) -> bool:
        if self.page_count <= 1:
            return False
        return (
            self.last_page_substantive_lines < ORPHAN_LAST_PAGE_MIN_LINES
            or self.last_page_fill_ratio < ORPHAN_LAST_PAGE_MIN_FILL_RATIO
        )


@dataclass(frozen=True)
class PaginationAdjustment:
    mode: str
    before: PaginationEstimate
    after: PaginationEstimate

    @property
    def adjusted(self) -> bool:
        return self.mode != "none"


def _length_points(value) -> float:
    if value is None:
        return 0.0
    if hasattr(value, "pt"):
        return float(value.pt)
    return 0.0


def _paragraph_format_value(paragraph, attribute: str):
    value = getattr(paragraph.paragraph_format, attribute)
    if value is not None:
        return value
    style = paragraph.style
    while style is not None:
        value = getattr(style.paragraph_format, attribute)
        if value is not None:
            return value
        style = style.base_style
    return None


def _paragraph_font_size(paragraph) -> float:
    sizes = [run.font.size.pt for run in paragraph.runs if run.font.size is not None]
    if sizes:
        return max(sizes)
    style = paragraph.style
    while style is not None:
        if style.font.size is not None:
            return float(style.font.size.pt)
        style = style.base_style
    return 10.0


def _line_spacing_points(paragraph, font_size: float) -> float:
    value = _paragraph_format_value(paragraph, "line_spacing")
    if hasattr(value, "pt"):
        return float(value.pt)
    if isinstance(value, (float, int)) and value > 0:
        return font_size * 1.2 * float(value)
    return font_size * 1.15


def _text_width_points(text: str, font_size: float) -> float:
    """Approximate Arial/Calibri-like text width without platform fonts.

    Resume export deliberately uses common sans-serif fonts. Character classes
    provide a more stable estimate than a raw characters-per-line rule and avoid
    a runtime dependency on installed font files.
    """
    em_units = 0.0
    for character in text:
        if character.isspace():
            em_units += 0.31
        elif character in "ilIjtfr.,:;|!'`":
            em_units += 0.31
        elif character in "MW@%&QO":
            em_units += 0.86
        elif character.isupper():
            em_units += 0.65
        elif character.isdigit():
            em_units += 0.58
        else:
            em_units += 0.54
    return em_units * font_size


def _paragraph_layouts(document) -> list[ParagraphLayoutEstimate]:
    if not document.sections:
        return []
    section = document.sections[0]
    page_width = section.page_width.inches if section.page_width is not None else 8.5
    left_margin = section.left_margin.inches if section.left_margin is not None else 0.5
    right_margin = section.right_margin.inches if section.right_margin is not None else 0.5
    usable_width_points = max(4.0 * 72.0, (page_width - left_margin - right_margin) * 72.0)

    rows: list[ParagraphLayoutEstimate] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        font_size = _paragraph_font_size(paragraph)
        line_spacing = _line_spacing_points(paragraph, font_size)
        before = _length_points(_paragraph_format_value(paragraph, "space_before"))
        after = _length_points(_paragraph_format_value(paragraph, "space_after"))
        left_indent = _length_points(_paragraph_format_value(paragraph, "left_indent"))
        right_indent = _length_points(_paragraph_format_value(paragraph, "right_indent"))
        available_width = max(144.0, usable_width_points - left_indent - right_indent)

        if text:
            if "\t" in paragraph.text:
                parts = paragraph.text.split("\t")
                text_width = sum(_text_width_points(part, font_size) for part in parts)
                text_width += max(0, len(parts) - 1) * 18.0
            else:
                text_width = _text_width_points(text, font_size)
            lines = max(1, ceil(text_width / available_width))
        else:
            lines = 0

        rows.append(
            ParagraphLayoutEstimate(
                index=index,
                height_points=(lines * line_spacing) + before + after,
                substantive_lines=lines,
                keep_with_next=_paragraph_format_value(paragraph, "keep_with_next") is True,
                page_break_before=_paragraph_format_value(paragraph, "page_break_before") is True,
                style_name=paragraph.style.name if paragraph.style is not None else "",
            )
        )
    return rows


def _keep_groups(rows: Iterable[ParagraphLayoutEstimate]) -> list[list[ParagraphLayoutEstimate]]:
    groups: list[list[ParagraphLayoutEstimate]] = []
    current: list[ParagraphLayoutEstimate] = []
    for row in rows:
        current.append(row)
        if not row.keep_with_next:
            groups.append(current)
            current = []
    if current:
        groups.append(current)
    return groups


def estimate_resume_pagination(document) -> PaginationEstimate:
    if not document.sections:
        return PaginationEstimate(1, (0.0,), (0,), 1.0)

    section = document.sections[0]
    page_height = section.page_height.inches if section.page_height is not None else 11.0
    top_margin = section.top_margin.inches if section.top_margin is not None else 0.5
    bottom_margin = section.bottom_margin.inches if section.bottom_margin is not None else 0.5
    usable_height = max(
        4.0 * 72.0,
        ((page_height - top_margin - bottom_margin) * 72.0)
        - PAGE_LAYOUT_SAFETY_RESERVE_POINTS,
    )

    used = [0.0]
    lines = [0]
    for group in _keep_groups(_paragraph_layouts(document)):
        if not group:
            continue
        if group[0].page_break_before and used[-1] > 0:
            used.append(0.0)
            lines.append(0)

        group_height = sum(row.height_points for row in group)
        if group_height <= usable_height and used[-1] > 0 and used[-1] + group_height > usable_height:
            used.append(0.0)
            lines.append(0)

        for row in group:
            if row.page_break_before and used[-1] > 0:
                used.append(0.0)
                lines.append(0)
            if used[-1] > 0 and used[-1] + row.height_points > usable_height:
                used.append(0.0)
                lines.append(0)
            used[-1] += row.height_points
            lines[-1] += row.substantive_lines

    return PaginationEstimate(
        page_count=max(1, len(used)),
        used_points_by_page=tuple(round(value, 3) for value in used),
        lines_by_page=tuple(lines),
        usable_height_points=usable_height,
    )


def _reduce_style_spacing(
    document,
    style_name: str,
    *,
    before_delta: float = 0.0,
    after_delta: float = 0.0,
    minimum_before: float = 0.0,
    minimum_after: float = 0.0,
) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    paragraph_format = style.paragraph_format
    if before_delta and paragraph_format.space_before is not None:
        paragraph_format.space_before = Pt(
            max(minimum_before, paragraph_format.space_before.pt - before_delta)
        )
    if after_delta and paragraph_format.space_after is not None:
        paragraph_format.space_after = Pt(
            max(minimum_after, paragraph_format.space_after.pt - after_delta)
        )


def _reduce_employer_gaps(document, amount: float, minimum: float) -> None:
    for paragraph in document.paragraphs:
        if paragraph.style is None or paragraph.style.name != STYLE_EMPLOYER_LINE:
            continue
        before = paragraph.paragraph_format.space_before
        if before is None or before.pt <= 0:
            continue
        paragraph.paragraph_format.space_before = Pt(max(minimum, before.pt - amount))


def _apply_compact_spacing(document, *, stronger: bool = False) -> None:
    # Keep 10 pt body text and standard 0.5 inch margins. The adjustment uses
    # whitespace that is visually expendable before touching readability.
    _reduce_style_spacing(
        document,
        STYLE_SECTION_HEADING,
        before_delta=0.5 if stronger else 1.5,
        after_delta=0.25 if stronger else 0.5,
        minimum_before=4.0 if stronger else 4.5,
        minimum_after=1.5 if stronger else 2.0,
    )
    _reduce_style_spacing(
        document,
        STYLE_BULLET,
        after_delta=0.2 if stronger else 0.3,
        minimum_after=0.0 if stronger else 0.2,
    )
    _reduce_style_spacing(
        document,
        STYLE_SUMMARY,
        after_delta=0.5,
        minimum_after=1.0 if stronger else 1.5,
    )
    _reduce_style_spacing(
        document,
        STYLE_SKILL_LINE,
        after_delta=0.25,
        minimum_after=0.0,
    )
    _reduce_style_spacing(
        document,
        STYLE_EDUCATION_DETAIL,
        after_delta=0.5,
        minimum_after=0.0,
    )
    _reduce_style_spacing(
        document,
        STYLE_EDUCATION,
        before_delta=0.5 if stronger else 0.0,
        minimum_before=0.5,
    )
    _reduce_style_spacing(
        document,
        STYLE_ROLE,
        after_delta=0.25 if stronger else 0.0,
        minimum_after=0.0,
    )
    for style_name in (STYLE_NAME, STYLE_TARGET_TITLE, STYLE_CONTACT):
        _reduce_style_spacing(
            document,
            style_name,
            after_delta=0.5,
            minimum_after=0.5 if stronger else 1.0,
        )
    _reduce_employer_gaps(
        document,
        amount=0.5 if stronger else 1.0,
        minimum=3.5 if stronger else 4.0,
    )


def _apply_balanced_page_break(document, estimate: PaginationEstimate) -> bool:
    rows = _paragraph_layouts(document)
    if len(rows) < 2:
        return False

    total_height = sum(row.height_points for row in rows)
    target_trailing_height = estimate.usable_height_points * 0.38
    candidates: list[tuple[float, int]] = []
    cumulative = 0.0
    for row in rows:
        trailing = total_height - cumulative
        if (
            row.index > 0
            and row.style_name in {STYLE_EMPLOYER_LINE, STYLE_SECTION_HEADING}
            and cumulative >= estimate.usable_height_points * 0.45
            and trailing >= estimate.usable_height_points * 0.28
            and trailing <= estimate.usable_height_points * 0.58
        ):
            candidates.append((abs(trailing - target_trailing_height), row.index))
        cumulative += row.height_points

    if not candidates:
        return False
    _, paragraph_index = min(candidates)
    document.paragraphs[paragraph_index].paragraph_format.page_break_before = True
    return True


def rebalance_resume_pagination(document) -> PaginationAdjustment:
    """Prevent a nearly empty final page without shrinking resume text.

    The exporter first tightens only discretionary spacing. If a genuinely long
    resume still leaves an orphan page, it moves a logical employer/section block
    to create a meaningful second page rather than leaving one isolated line.
    """
    before = estimate_resume_pagination(document)
    if not before.has_orphan_final_page:
        return PaginationAdjustment("none", before, before)

    _apply_compact_spacing(document)
    after = estimate_resume_pagination(document)
    if not after.has_orphan_final_page:
        return PaginationAdjustment("compact", before, after)

    _apply_compact_spacing(document, stronger=True)
    after = estimate_resume_pagination(document)
    if not after.has_orphan_final_page:
        return PaginationAdjustment("extra_compact", before, after)

    if _apply_balanced_page_break(document, after):
        balanced = estimate_resume_pagination(document)
        return PaginationAdjustment("balanced", before, balanced)

    return PaginationAdjustment("extra_compact", before, after)
