from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .bullet_text import (
    has_bullet_structure_artifacts,
    normalize_resume_bullet_terminal_punctuation,
    normalize_resume_bullet_text,
)
from .docx_styles import (
    DEFAULT_RESUME_STYLE,
    RESUME_STYLE_THEMES,
    ResumeStyleTheme,
    STYLE_BULLET,
    STYLE_CONTACT,
    STYLE_EDUCATION,
    STYLE_EDUCATION_DETAIL,
    STYLE_EDUCATION_META,
    STYLE_EMPLOYER_LINE,
    STYLE_NAME,
    STYLE_ROLE,
    STYLE_SECTION_HEADING,
    STYLE_SKILL_LINE,
    STYLE_SUMMARY,
    STYLE_TARGET_TITLE,
    clear_document_body,
    clear_headers_and_footers,
    configure_resume_document,
    normalize_career_stage,
    normalize_resume_format,
    normalize_visual_design,
    resume_preference_label,
)
from .models import ApprovedResume, CandidateProfile
from .resume_pagination import rebalance_resume_pagination
from .resume_language import resume_format_headings, resume_labels
from .validation import adjacent_repeated_words


class TemplateError(RuntimeError):
    """Raised when the Word template cannot produce a valid resume document."""


RESUME_FORMAT_SECTIONS = {
    "standard": {
        "summary": "Professional Summary",
        "skills": "Skills",
        "experience": "Work Experience",
        "education": "Education and Professional Development",
    },
    "technical": {
        "summary": "Technical Profile",
        "skills": "Technical Skills",
        "experience": "Engineering Experience",
        "education": "Education and Certifications",
    },
    "career_changer": {
        "summary": "Professional Profile",
        "skills": "Transferable and Relevant Skills",
        "experience": "Relevant Experience",
        "education": "Education and Professional Development",
    },
    "freelance": {
        "summary": "Professional Profile",
        "skills": "Core Capabilities",
        "experience": "Client and Project Experience",
        "education": "Education and Credentials",
    },
}

SKILL_CATEGORY_ORDER = {
    "standard": ("hard", "soft", "tools", "industry", "languages"),
    "technical": ("tools", "hard", "industry", "soft", "languages"),
    "career_changer": ("soft", "hard", "industry", "tools", "languages"),
    "freelance": ("hard", "tools", "industry", "soft", "languages"),
}


def _document_text_blocks(document) -> list[str]:
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.extend(
                    paragraph.text
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
    return blocks


def _enforce_repeated_word_export_gate(document) -> None:
    repeated: list[str] = []
    seen: set[str] = set()
    for block in _document_text_blocks(document):
        for word in adjacent_repeated_words(block):
            key = word.casefold()
            if key not in seen:
                repeated.append(word)
                seen.add(key)
    if repeated:
        pairs = ", ".join(f"'{word} {word}'" for word in repeated)
        raise TemplateError(
            "Export blocked because adjacent repeated word(s) remain in the generated resume: "
            + pairs
            + ". Correct the highlighted resume text and run verification again."
        )


def _add_hyperlink(
    paragraph, text: str, url: str, theme: ResumeStyleTheme
) -> None:
    """Append an external hyperlink without relying on hardcoded template relationships."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), theme.body_font)
    run_properties.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), theme.link_color)
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(round(theme.contact_size * 2))))
    run_properties.append(size)
    size_complex = OxmlElement("w:szCs")
    size_complex.set(qn("w:val"), str(int(round(theme.contact_size * 2))))
    run_properties.append(size_complex)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _usable_width_inches(document) -> float:
    section = document.sections[0]
    width = section.page_width.inches if section.page_width is not None else 8.5
    left = section.left_margin.inches if section.left_margin is not None else 0.6
    right = section.right_margin.inches if section.right_margin is not None else 0.6
    return max(4.0, width - left - right)


def _add_right_tab_stop(document, paragraph) -> None:
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(_usable_width_inches(document)),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.SPACES,
    )


def _add_section_heading(
    document, text: str, theme: ResumeStyleTheme
) -> None:
    label = text.upper() if theme.section_uppercase else text
    document.add_paragraph(label, style=STYLE_SECTION_HEADING)


def _append_contact_items(
    paragraph,
    items: list[tuple[str, str]],
    theme: ResumeStyleTheme,
) -> None:
    for index, (label, target) in enumerate(items):
        if index:
            paragraph.add_run("  |  ")
        if target:
            _add_hyperlink(paragraph, label, target, theme)
        else:
            paragraph.add_run(label)


def _add_header(
    document, profile: CandidateProfile, approved: ApprovedResume, theme: ResumeStyleTheme
) -> None:
    document.add_paragraph(profile.name.strip(), style=STYLE_NAME)
    if approved.target_title.strip():
        document.add_paragraph(approved.target_title.strip(), style=STYLE_TARGET_TITLE)

    contact = profile.contact
    linked_items: list[tuple[str, str]] = []
    if contact.phone.strip():
        linked_items.append((contact.phone.strip(), ""))
    if contact.email.strip():
        linked_items.append((contact.email.strip(), f"mailto:{contact.email.strip()}"))
    if contact.linkedin_url.strip():
        linked_items.append((contact.linkedin_label.strip() or "LinkedIn", contact.linkedin_url.strip()))
    if contact.github_url.strip():
        linked_items.append((contact.github_label.strip() or "GitHub", contact.github_url.strip()))

    if theme.is_mid_career_corporate:
        # Match the attached original resume: location on its own centered line,
        # followed by the phone/email/profile links on a second centered line.
        if contact.location.strip():
            location = document.add_paragraph(contact.location.strip(), style=STYLE_CONTACT)
            location.paragraph_format.space_after = Inches(0.01)
        if linked_items:
            paragraph = document.add_paragraph(style=STYLE_CONTACT)
            _append_contact_items(paragraph, linked_items, theme)
        return

    contact_items = list(linked_items)
    if contact.location.strip():
        contact_items.insert(0, (contact.location.strip(), ""))
    if contact_items:
        paragraph = document.add_paragraph(style=STYLE_CONTACT)
        _append_contact_items(paragraph, contact_items, theme)


def _add_summary(
    document,
    approved: ApprovedResume,
    theme: ResumeStyleTheme,
    *,
    heading: str = "Professional Summary",
) -> None:
    summary = approved.professional_summary.strip()
    if not summary:
        return
    _add_section_heading(document, heading, theme)
    document.add_paragraph(summary, style=STYLE_SUMMARY)


def _add_skills(
    document,
    profile: CandidateProfile,
    approved: ApprovedResume,
    theme: ResumeStyleTheme,
    *,
    heading: str = "Skills",
    resume_format: str = "standard",
    resume_language: str = "English",
) -> None:
    labels = resume_labels(resume_language)
    category_map = {
        "hard": (labels["hard_skills"], approved.skills.hard_skills),
        "soft": (labels["soft_skills"], approved.skills.soft_skills),
        "tools": (labels["tools_software"], approved.skills.tools_software),
        "industry": (labels["industry_knowledge"], approved.skills.industry_knowledge),
        "languages": (labels["languages"], profile.skills.languages),
    }
    sections = [category_map[key] for key in SKILL_CATEGORY_ORDER[resume_format]]
    populated = [(label, [item.strip() for item in items if item.strip()]) for label, items in sections]
    populated = [(label, items) for label, items in populated if items]
    if not populated:
        return

    _add_section_heading(document, heading, theme)
    if theme.is_mid_career_corporate:
        # The original resume uses one compact skills block with line breaks.
        paragraph = document.add_paragraph(style=STYLE_SKILL_LINE)
        for index, (label, items) in enumerate(populated):
            if index:
                paragraph.add_run().add_break()
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            paragraph.add_run(", ".join(items))
        return

    for label, items in populated:
        paragraph = document.add_paragraph(style=STYLE_SKILL_LINE)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        if (theme.is_executive or theme.is_early_career):
            label_run.font.color.rgb = theme.accent_color
        paragraph.add_run(", ".join(items))


def _add_experience(
    document,
    profile: CandidateProfile,
    approved: ApprovedResume,
    theme: ResumeStyleTheme,
    *,
    heading: str = "Work Experience",
) -> None:
    included = [
        (experience, [item.strip() for item in approved.bullets_by_experience.get(experience.id, []) if item.strip()])
        for experience in profile.experiences
    ]
    included = [(experience, bullets) for experience, bullets in included if bullets]
    if not included:
        return

    _add_section_heading(document, heading, theme)
    for index, (experience, bullets) in enumerate(included):
        employer_line = document.add_paragraph(style=STYLE_EMPLOYER_LINE)
        # Let Word paginate naturally. Employer headings are already configured
        # to stay with their role title, so a forced break is unnecessary and
        # can leave a largely empty second page.
        _add_right_tab_stop(document, employer_line)
        employer_line.paragraph_format.space_before = Pt(
            0 if index == 0 else 6 if theme.is_mid_career_corporate else 4.5 if theme.is_executive else 4
        )
        employer_run = employer_line.add_run(experience.employer.strip())
        employer_run.bold = True
        employer_run.underline = False
        location = experience.location.strip()
        if location:
            employer_line.add_run(f", {location}")
        if experience.dates.strip():
            employer_line.add_run("\t" + experience.dates.strip())

        document.add_paragraph(experience.title.strip(), style=STYLE_ROLE)
        for bullet in bullets:
            rendered_bullet = (
                normalize_resume_bullet_text(bullet, max_words=35)
                if has_bullet_structure_artifacts(bullet)
                else " ".join(bullet.split()).strip()
            )
            rendered_bullet = normalize_resume_bullet_terminal_punctuation(rendered_bullet)
            if not rendered_bullet:
                continue
            paragraph = document.add_paragraph(style=STYLE_BULLET)
            # The export owns the bullet glyph. proposed_text must remain plain text
            # so nested markdown cannot create duplicate or malformed list markers.
            marker = paragraph.add_run("• ")
            if theme.is_early_career:
                marker.font.color.rgb = theme.accent_color
                marker.bold = True
            paragraph.add_run(rendered_bullet)


def _add_education(
    document,
    profile: CandidateProfile,
    theme: ResumeStyleTheme,
    *,
    heading: str = "Education and Professional Development",
) -> None:
    if not profile.education:
        return
    _add_section_heading(document, heading, theme)
    for item in profile.education:
        paragraph = document.add_paragraph(style=STYLE_EDUCATION)
        _add_right_tab_stop(document, paragraph)
        credential = item.credential.strip()
        if credential:
            run = paragraph.add_run(credential)
            run.bold = True

        institution_parts = [part for part in (item.institution.strip(), item.location.strip()) if part]
        if theme.is_mid_career_corporate and institution_parts:
            if credential:
                paragraph.add_run(", ")
            paragraph.add_run(", ".join(institution_parts))
        if item.date.strip():
            paragraph.add_run("\t" + item.date.strip())

        if theme.is_mid_career_corporate:
            paragraph.paragraph_format.keep_with_next = bool(item.detail.strip())
            if item.detail.strip():
                detail = document.add_paragraph(style=STYLE_EDUCATION_DETAIL)
                detail.add_run("• ")
                detail.add_run(
                    normalize_resume_bullet_terminal_punctuation(item.detail.strip())
                )
            continue

        has_following_line = bool(institution_parts or item.detail.strip())
        paragraph.paragraph_format.keep_with_next = has_following_line
        if institution_parts:
            institution = document.add_paragraph(", ".join(institution_parts), style=STYLE_EDUCATION_META)
            institution.paragraph_format.keep_with_next = bool(item.detail.strip())
        if item.detail.strip():
            document.add_paragraph(
                normalize_resume_bullet_terminal_punctuation(item.detail.strip()),
                style=STYLE_EDUCATION_DETAIL,
            )


def _validate_generated_document(document) -> None:
    text = "\n".join(_document_text_blocks(document))
    if not text.strip():
        raise TemplateError("The generated resume is empty.")
    if "{{" in text or "}}" in text:
        raise TemplateError("The generated resume contains an unresolved template marker.")


def export_resume_docx(
    template_path: str | Path,
    profile: CandidateProfile,
    approved: ApprovedResume,
    *,
    enforce_language_gate: bool = True,
    style_key: str | None = None,
    career_stage: str | None = None,
    resume_format: str | None = None,
    visual_design: str | None = None,
    resume_language: str | None = None,
) -> bytes:
    template = Path(template_path)
    if not template.exists():
        raise TemplateError(f"Resume template was not found: {template}")
    try:
        document = Document(str(template))
    except Exception as exc:
        raise TemplateError(f"Resume template could not be opened: {exc}") from exc

    if career_stage is None:
        effective_style = style_key
        if effective_style is None:
            template_name = template.stem.casefold()
            effective_style = next(
                (key for key in RESUME_STYLE_THEMES if key in template_name),
                DEFAULT_RESUME_STYLE,
            )
        career_stage = normalize_career_stage(effective_style)
    stage = normalize_career_stage(career_stage)
    format_key = normalize_resume_format(resume_format)
    design_key = normalize_visual_design(visual_design)
    theme = configure_resume_document(
        document,
        style_key or DEFAULT_RESUME_STYLE,
        career_stage=stage,
        visual_design=design_key,
    )
    clear_document_body(document)
    clear_headers_and_footers(document)

    headings = resume_format_headings(resume_language or "English", format_key)
    _add_header(document, profile, approved, theme)

    def add_summary() -> None:
        _add_summary(document, approved, theme, heading=headings["summary"])

    def add_skills() -> None:
        _add_skills(
            document,
            profile,
            approved,
            theme,
            heading=headings["skills"],
            resume_format=format_key,
            resume_language=resume_language or "English",
        )

    def add_experience() -> None:
        _add_experience(
            document, profile, approved, theme, heading=headings["experience"]
        )

    def add_education() -> None:
        _add_education(document, profile, theme, heading=headings["education"])

    if format_key == "technical":
        add_skills()
        add_summary()
        add_experience()
        add_education()
    elif format_key == "career_changer":
        add_summary()
        add_skills()
        add_education()
        add_experience()
    elif format_key == "freelance":
        add_summary()
        add_skills()
        add_experience()
        add_education()
    else:
        add_summary()
        add_skills()
        if theme.is_early_career:
            add_education()
            add_experience()
        else:
            add_experience()
            add_education()

    preference_label = resume_preference_label(stage, format_key, design_key)
    document.core_properties.title = f"{profile.name} - {approved.target_title or 'Resume'}"
    document.core_properties.subject = f"Tailored resume - {preference_label}"
    document.core_properties.author = profile.name
    document.core_properties.keywords = (
        f"resume, ATS, tailored, {stage}, {format_key}, {design_key}"
    )

    rebalance_resume_pagination(document)
    _validate_generated_document(document)
    if enforce_language_gate:
        _enforce_repeated_word_export_gate(document)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
