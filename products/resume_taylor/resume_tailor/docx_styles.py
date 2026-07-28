from __future__ import annotations

from dataclasses import dataclass, replace
from datetime import date
import re
from typing import Any, Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

STYLE_NAME = "Resume Name"
STYLE_TARGET_TITLE = "Resume Target Title"
STYLE_CONTACT = "Resume Contact"
STYLE_SECTION_HEADING = "Resume Section Heading"
STYLE_SUMMARY = "Resume Summary"
STYLE_SKILL_LINE = "Resume Skill Line"
STYLE_EMPLOYER_LINE = "Resume Employer and Dates"
STYLE_ROLE = "Resume Job Title"
STYLE_BULLET = "Resume Experience Bullet"
STYLE_EDUCATION = "Resume Education Entry"
STYLE_EDUCATION_META = "Resume Education Institution"
STYLE_EDUCATION_DETAIL = "Resume Education Detail"

RESUME_STYLE_NAMES = (
    STYLE_NAME,
    STYLE_TARGET_TITLE,
    STYLE_CONTACT,
    STYLE_SECTION_HEADING,
    STYLE_SUMMARY,
    STYLE_SKILL_LINE,
    STYLE_EMPLOYER_LINE,
    STYLE_ROLE,
    STYLE_BULLET,
    STYLE_EDUCATION,
    STYLE_EDUCATION_META,
    STYLE_EDUCATION_DETAIL,
)


@dataclass(frozen=True)
class ResumeStyleTheme:
    key: str
    label: str
    audience: str
    description: str
    body_font: str
    heading_font: str
    accent_color: RGBColor
    text_color: RGBColor
    link_color: str
    header_alignment: WD_ALIGN_PARAGRAPH
    section_uppercase: bool
    section_border: bool
    top_margin: float
    bottom_margin: float
    left_margin: float
    right_margin: float
    name_size: float
    target_size: float
    contact_size: float
    section_size: float
    body_size: float
    skill_size: float
    employer_size: float
    role_size: float
    bullet_size: float
    education_size: float
    section_space_before: float
    section_space_after: float
    bullet_space_after: float
    collection: str = "career_stage"
    selector_note: str = ""
    career_stage: str = ""
    visual_design: str = "corporate"

    @property
    def stage_key(self) -> str:
        if self.career_stage:
            return self.career_stage
        return {
            "early_career": "early_career",
            "professional": "mid_career",
            "executive": "executive",
        }.get(self.key, "mid_career")

    @property
    def is_early_career(self) -> bool:
        return self.stage_key == "early_career"

    @property
    def is_mid_career(self) -> bool:
        return self.stage_key == "mid_career"

    @property
    def is_executive(self) -> bool:
        return self.stage_key == "executive"

    @property
    def is_corporate(self) -> bool:
        return self.visual_design == "corporate"

    @property
    def is_modern(self) -> bool:
        return self.visual_design == "modern"

    @property
    def is_mid_career_corporate(self) -> bool:
        return self.is_mid_career and self.is_corporate


EARLY_CAREER_STYLE = ResumeStyleTheme(
    key="early_career",
    label="Early Career",
    audience="Students, Graduates & Career Changers",
    description=(
        "Approachable and one-page friendly, with education placed before work "
        "experience and stronger emphasis on skills and transferable experience."
    ),
    body_font="Arial",
    heading_font="Arial",
    accent_color=RGBColor(37, 99, 235),
    text_color=RGBColor(30, 41, 59),
    link_color="2563EB",
    header_alignment=WD_ALIGN_PARAGRAPH.CENTER,
    section_uppercase=False,
    section_border=False,
    top_margin=0.58,
    bottom_margin=0.58,
    left_margin=0.68,
    right_margin=0.68,
    name_size=20,
    target_size=11.25,
    contact_size=9,
    section_size=10.25,
    body_size=9.4,
    skill_size=9.0,
    employer_size=9.5,
    role_size=9.5,
    bullet_size=9.35,
    education_size=9.3,
    section_space_before=6.5,
    section_space_after=2.5,
    bullet_space_after=1.0,
)


# This style intentionally recreates the visual language of the user's original
# resume: centered blue headings, expanded lettering, a ruled title block,
# compact Calibri typography, underlined employer names, and concise spacing.
PROFESSIONAL_STYLE = ResumeStyleTheme(
    key="professional",
    label="Mid-Career Professional",
    audience="Experienced Professionals",
    description=(
        "A polished, compact style based on your original resume, with centered "
        "blue headings and a familiar experience-first structure."
    ),
    body_font="Calibri",
    heading_font="Calibri",
    accent_color=RGBColor(0, 85, 161),
    text_color=RGBColor(0, 0, 0),
    link_color="2E687C",
    header_alignment=WD_ALIGN_PARAGRAPH.CENTER,
    section_uppercase=False,
    section_border=False,
    top_margin=0.5,
    bottom_margin=0.5,
    left_margin=0.5,
    right_margin=0.5,
    name_size=15,
    target_size=12,
    contact_size=10.5,
    section_size=12,
    body_size=10,
    skill_size=10,
    employer_size=10,
    role_size=10,
    bullet_size=10,
    education_size=10,
    section_space_before=7,
    section_space_after=3,
    bullet_space_after=0.5,
    selector_note="Recreates the centered blue layout and compact typography of your original resume.",
)


EXECUTIVE_STYLE = ResumeStyleTheme(
    key="executive",
    label="Executive Leadership",
    audience="Directors, VPs & Senior Leaders",
    description=(
        "Authoritative, achievement-led, and spacious for directors, executives, "
        "transformation leaders, and senior consultants."
    ),
    body_font="Arial",
    heading_font="Cambria",
    accent_color=RGBColor(37, 55, 82),
    text_color=RGBColor(17, 24, 39),
    link_color="253752",
    header_alignment=WD_ALIGN_PARAGRAPH.LEFT,
    section_uppercase=False,
    section_border=True,
    top_margin=0.55,
    bottom_margin=0.55,
    left_margin=0.65,
    right_margin=0.65,
    name_size=22,
    target_size=12,
    contact_size=9,
    section_size=10.75,
    body_size=9.5,
    skill_size=9.1,
    employer_size=9.8,
    role_size=9.6,
    bullet_size=9.5,
    education_size=9.25,
    section_space_before=7,
    section_space_after=2.5,
    bullet_space_after=1.0,
)


# Keep the options in career-stage order in the selector.
RESUME_STYLE_THEMES = {
    EARLY_CAREER_STYLE.key: EARLY_CAREER_STYLE,
    PROFESSIONAL_STYLE.key: PROFESSIONAL_STYLE,
    EXECUTIVE_STYLE.key: EXECUTIVE_STYLE,
}
DEFAULT_RESUME_STYLE = PROFESSIONAL_STYLE.key
DEFAULT_CAREER_STAGE = "mid_career"
DEFAULT_RESUME_FORMAT = "standard"
DEFAULT_VISUAL_DESIGN = "corporate"

CAREER_STAGE_OPTIONS = (
    {
        "key": "early_career",
        "label": "Early Career",
        "audience": "Students, graduates, and professionals building experience",
        "description": "Keeps the document concise and gives education and transferable experience more prominence.",
    },
    {
        "key": "mid_career",
        "label": "Mid-Career Professional",
        "audience": "Experienced individual contributors and managers",
        "description": "Balances depth, measurable outcomes, skills, and a clear reverse-chronological career history.",
    },
    {
        "key": "executive",
        "label": "Executive Leadership",
        "audience": "Directors, vice presidents, and senior leaders",
        "description": "Creates stronger hierarchy and space for enterprise scope, leadership, and strategic accomplishments.",
    },
)

RESUME_FORMAT_OPTIONS = (
    {
        "key": "standard",
        "label": "Standard Professional",
        "audience": "Most traditional applications",
        "description": "Uses a familiar reverse-chronological structure with summary, skills, experience, and education.",
    },
    {
        "key": "technical",
        "label": "Technical / Engineering",
        "audience": "Software, data, IT, cloud, and engineering roles",
        "description": "Moves the technical stack forward and uses engineering-focused section labels and skill ordering.",
    },
    {
        "key": "career_changer",
        "label": "Career Changer / Hybrid",
        "audience": "Industry pivots, re-entry, and transferable experience",
        "description": "Leads with relevant and transferable capabilities before presenting the full employment chronology.",
    },
    {
        "key": "freelance",
        "label": "Freelance / Project-Based",
        "audience": "Consultants, contractors, and independent professionals",
        "description": "Frames employment entries as client and project experience, emphasizing deliverables and outcomes.",
    },
)

VISUAL_DESIGN_OPTIONS = (
    {
        "key": "corporate",
        "label": "Corporate",
        "audience": "Banking, finance, audit, government, and established employers",
        "description": "Conservative typography and compact spacing for a traditional, highly familiar Word resume.",
    },
    {
        "key": "modern",
        "label": "Modern",
        "audience": "Software, technology, product, and contemporary organizations",
        "description": "Clean left-aligned hierarchy, contemporary spacing, and restrained blue accents while remaining ATS-friendly.",
    },
)

_STAGE_TO_TEMPLATE_STYLE = {
    "early_career": "early_career",
    "mid_career": "professional",
    "executive": "executive",
}
_TEMPLATE_STYLE_TO_STAGE = {value: key for key, value in _STAGE_TO_TEMPLATE_STYLE.items()}


def normalize_career_stage(value: str | None) -> str:
    key = (value or "").strip().casefold()
    if key in _STAGE_TO_TEMPLATE_STYLE:
        return key
    return _TEMPLATE_STYLE_TO_STAGE.get(key, DEFAULT_CAREER_STAGE)


def career_stage_template_key(value: str | None) -> str:
    return _STAGE_TO_TEMPLATE_STYLE[normalize_career_stage(value)]


def normalize_resume_format(value: str | None) -> str:
    key = (value or "").strip().casefold()
    allowed = {option["key"] for option in RESUME_FORMAT_OPTIONS}
    return key if key in allowed else DEFAULT_RESUME_FORMAT


def normalize_visual_design(value: str | None) -> str:
    key = (value or "").strip().casefold()
    allowed = {option["key"] for option in VISUAL_DESIGN_OPTIONS}
    return key if key in allowed else DEFAULT_VISUAL_DESIGN


def normalize_resume_style(value: str | None) -> str:
    """Normalize the legacy one-dimensional style to a template key."""
    return career_stage_template_key(value)


def get_resume_style_theme(value: str | None) -> ResumeStyleTheme:
    return RESUME_STYLE_THEMES[normalize_resume_style(value)]


def compose_resume_theme(
    career_stage: str | None,
    visual_design: str | None,
) -> ResumeStyleTheme:
    stage = normalize_career_stage(career_stage)
    design = normalize_visual_design(visual_design)
    base = RESUME_STYLE_THEMES[career_stage_template_key(stage)]
    if design == "corporate":
        return replace(base, career_stage=stage, visual_design=design)

    # The modern design deliberately changes visual presentation only. Career
    # stage continues to control spacing, density, and content hierarchy.
    return replace(
        base,
        career_stage=stage,
        visual_design=design,
        body_font="Arial",
        heading_font="Arial",
        accent_color=RGBColor(37, 99, 235),
        text_color=RGBColor(30, 41, 59),
        link_color="2563EB",
        header_alignment=WD_ALIGN_PARAGRAPH.LEFT,
        section_uppercase=False,
        section_border=False,
        name_size=max(base.name_size, 19 if stage != "executive" else 22),
        target_size=max(base.target_size, 11.25),
        contact_size=min(base.contact_size, 9.25),
        section_size=max(base.section_size, 10.25),
        section_space_before=max(base.section_space_before, 6),
        section_space_after=max(base.section_space_after, 2.5),
    )


def _candidate_answer_text(candidate_answers: Iterable[Any] | None) -> str:
    """Return candidate-confirmed evidence as searchable text.

    Answers are intentionally treated as supplemental evidence only. A bare
    yes/no response does not create expertise; only the question and any
    candidate-provided details contribute context.
    """
    parts: list[str] = []
    for answer in candidate_answers or ():
        question = str(getattr(answer, "question", "") or "").strip()
        text = str(getattr(answer, "text", "") or "").strip()
        if text:
            parts.extend((question, text))
    return "\n".join(part for part in parts if part)


def _candidate_context_text(
    candidate_profile: Any | None,
    candidate_answers: Iterable[Any] | None = None,
) -> str:
    parts: list[str] = []
    if candidate_profile is not None:
        all_source_text = getattr(candidate_profile, "all_source_text", None)
        if callable(all_source_text):
            parts.append(str(all_source_text()))
        else:
            parts.append(str(candidate_profile))
    answer_text = _candidate_answer_text(candidate_answers)
    if answer_text:
        parts.append(answer_text)
    return "\n".join(parts).casefold()


def _candidate_experience_years(candidate_profile: Any | None) -> float | None:
    """Estimate non-overlapping experience years from summary and date ranges."""
    if candidate_profile is None:
        return None

    summary = str(getattr(candidate_profile, "current_summary", "") or "")
    explicit_years = [
        int(value)
        for value in re.findall(
            r"\b(?:over\s+|more\s+than\s+|about\s+|approximately\s+|nearly\s+)?(\d{1,2})\+?\s+years?\b",
            summary,
            flags=re.IGNORECASE,
        )
        if 0 <= int(value) <= 50
    ]

    intervals: list[tuple[int, int]] = []
    current_month = date.today().year * 12 + date.today().month - 1
    for experience in getattr(candidate_profile, "experiences", ()) or ():
        raw_dates = str(getattr(experience, "dates", "") or "")
        matches = re.findall(r"(?:(\d{1,2})[/-])?((?:19|20)\d{2})", raw_dates)
        if not matches:
            continue
        start_month_text, start_year_text = matches[0]
        start_month = int(start_month_text or 1)
        start = int(start_year_text) * 12 + max(1, min(start_month, 12)) - 1
        if len(matches) >= 2:
            end_month_text, end_year_text = matches[-1]
            end_month = int(end_month_text or 12)
            end = int(end_year_text) * 12 + max(1, min(end_month, 12)) - 1
        elif re.search(r"\b(?:present|current|now)\b", raw_dates, re.IGNORECASE):
            end = current_month
        else:
            continue
        if end >= start:
            intervals.append((start, end + 1))

    merged: list[list[int]] = []
    for start, end in sorted(intervals):
        if not merged or start > merged[-1][1]:
            merged.append([start, end])
        else:
            merged[-1][1] = max(merged[-1][1], end)
    merged_months = sum(end - start for start, end in merged)

    date_years = merged_months / 12 if merged_months else None
    if explicit_years and date_years is not None:
        return max(float(max(explicit_years)), date_years)
    if explicit_years:
        return float(max(explicit_years))
    return date_years


def _term_score(text: str, terms: Iterable[str]) -> int:
    return sum(1 for term in terms if term in text)


_DOMAIN_TERMS: dict[str, tuple[str, ...]] = {
    "technical": (
        "software", "developer", "engineering", "engineer", "programming",
        "python", "java", "javascript", "sql", "database", "cloud", "aws",
        "azure", "devops", "sre", "cybersecurity", "machine learning",
        "data science", "data engineer", "api", "full stack", "backend",
        "frontend", "quality assurance", "qa engineer",
    ),
    "finance_risk": (
        "banking", "bank", "financial", "finance", "accounting", "audit",
        "auditor", "regulatory", "compliance", "insurance", "risk",
        "internal controls", "tax", "treasury", "investment",
    ),
    "sales_marketing": (
        "sales", "business development", "account executive", "marketing",
        "brand", "campaign", "seo", "content marketing", "growth marketing",
        "customer acquisition", "demand generation", "public relations",
    ),
    "operations_project": (
        "operations", "supply chain", "logistics", "program manager",
        "project manager", "project management", "process improvement",
        "procurement", "vendor management", "service delivery",
    ),
    "people_hr": (
        "human resources", "hr business partner", "recruiting", "recruiter",
        "talent acquisition", "people operations", "employee relations",
        "compensation", "learning and development",
    ),
    "healthcare": (
        "healthcare", "clinical", "patient", "hospital", "medical", "nursing",
        "pharmaceutical", "public health", "health system",
    ),
    "legal": (
        "attorney", "lawyer", "legal counsel", "litigation", "paralegal",
        "contract law", "legal research", "case management",
    ),
    "education_research": (
        "teacher", "teaching", "curriculum", "professor", "academic",
        "researcher", "research scientist", "laboratory", "publication",
        "higher education", "instructional design",
    ),
    "creative_design": (
        "graphic design", "ux designer", "ui designer", "creative director",
        "art director", "copywriter", "portfolio", "visual design",
    ),
}


def _domain_scores(text: str) -> dict[str, int]:
    return {name: _term_score(text, terms) for name, terms in _DOMAIN_TERMS.items()}


def _candidate_is_established(candidate_profile: Any | None) -> bool:
    years = _candidate_experience_years(candidate_profile)
    experience_count = len(getattr(candidate_profile, "experiences", ()) or ())
    return bool((years is not None and years >= 4) or experience_count >= 2)


def _candidate_freelance_score(candidate_profile: Any | None, candidate_text: str) -> int:
    engagement_terms = (
        "freelance", "independent consultant", "independent contractor",
        "self-employed", "self employed", "fractional", "contract consultant",
    )
    score = _term_score(candidate_text, engagement_terms)
    engagement_count = 0
    for experience in getattr(candidate_profile, "experiences", ()) or ():
        heading = " ".join(
            str(getattr(experience, field, "") or "")
            for field in ("title", "employer")
        ).casefold()
        if any(term in heading for term in engagement_terms):
            engagement_count += 1
    if engagement_count >= 2:
        score += 3
    elif engagement_count == 1:
        score += 1
    return score


def _looks_like_career_change(
    job_text: str,
    candidate_profile: Any | None,
    candidate_text: str,
) -> bool:
    explicit_change_terms = (
        "career change", "career changer", "changing careers", "transitioning into",
        "pivoting into", "returning to the workforce", "workforce re-entry",
    )
    if any(term in candidate_text for term in explicit_change_terms):
        return True
    if not _candidate_is_established(candidate_profile):
        return False

    job_scores = _domain_scores(job_text)
    candidate_scores = _domain_scores(candidate_text)
    strongest_job_domain, strongest_job_score = max(
        job_scores.items(), key=lambda item: item[1]
    )
    if strongest_job_score < 2:
        return False

    target_candidate_score = candidate_scores[strongest_job_domain]
    strongest_candidate_score = max(candidate_scores.values(), default=0)
    return target_candidate_score == 0 and strongest_candidate_score >= 2


def recommend_career_stage(
    job_description: str,
    target_title: str = "",
    *,
    candidate_profile: Any | None = None,
    candidate_answers: Iterable[Any] | None = None,
) -> str:
    """Recommend career stage from both the target role and candidate evidence."""
    title = target_title.casefold()
    job_text = f"{target_title} {job_description}".casefold()
    candidate_text = _candidate_context_text(candidate_profile, candidate_answers)
    candidate_years = _candidate_experience_years(candidate_profile)
    has_candidate_context = candidate_profile is not None or bool(candidate_text)

    executive_title_terms = (
        "chief", "ceo", "cto", "cfo", "coo", "vice president", "vp",
        "director", "head of", "general manager", "executive",
    )
    executive_context_terms = (
        "board", "enterprise strategy", "organizational transformation",
        "executive leadership", "p&l", "portfolio strategy", "business unit",
    )
    leadership_evidence_terms = (
        "led ", "leadership", "managed a team", "managed teams", "people manager",
        "direct reports", "enterprise-wide", "organization-wide", "strategic roadmap",
        "executive communication", "senior management",
    )
    senior_candidate_title_terms = (
        "director", "vice president", "vp", "head of", "chief", "executive",
        "general manager", "senior manager",
    )
    early_title_terms = (
        "intern", "internship", "entry level", "entry-level", "junior",
        "new graduate", "new grad", "graduate trainee", "apprentice",
    )
    early_context_terms = (
        "recent graduate", "current student", "campus recruiting",
        "graduate program", "no prior experience", "early career",
    )

    executive_score = (
        2 * _term_score(title, executive_title_terms)
        + _term_score(job_text, executive_context_terms)
    )
    early_score = (
        2 * _term_score(title, early_title_terms)
        + _term_score(job_text, early_context_terms)
    )
    candidate_exec_titles = _term_score(candidate_text, senior_candidate_title_terms)
    candidate_leadership = _term_score(candidate_text, leadership_evidence_terms)
    candidate_has_senior_scope = bool(
        candidate_exec_titles >= 1
        or (
            candidate_years is not None
            and candidate_years >= 10
            and candidate_leadership >= 2
        )
    )

    if executive_score >= 2:
        if not has_candidate_context:
            return "executive"
        # Avoid presenting an unsupported executive document solely because the
        # desired title contains "Director" or "VP".
        return "executive" if candidate_has_senior_scope else "mid_career"

    if early_score >= 2:
        if not has_candidate_context:
            return "early_career"
        # An experienced applicant targeting a junior role still needs a mature
        # resume rather than a graduate-style document.
        if candidate_years is None:
            return "early_career"
        return "early_career" if candidate_years < 4 and not candidate_has_senior_scope else "mid_career"

    if candidate_years is not None and candidate_years < 3 and not candidate_has_senior_scope:
        return "early_career"

    if not target_title.strip() and candidate_exec_titles >= 1 and candidate_leadership >= 2:
        return "executive"

    return "mid_career"


def recommend_resume_style(
    job_description: str,
    target_title: str = "",
    *,
    candidate_profile: Any | None = None,
    candidate_answers: Iterable[Any] | None = None,
) -> str:
    """Backward-compatible recommendation returning a template style key."""
    return career_stage_template_key(
        recommend_career_stage(
            job_description,
            target_title,
            candidate_profile=candidate_profile,
            candidate_answers=candidate_answers,
        )
    )


def recommend_resume_format(
    job_description: str,
    target_title: str = "",
    *,
    candidate_profile: Any | None = None,
    candidate_answers: Iterable[Any] | None = None,
) -> str:
    """Recommend structure from role needs and the candidate's career shape."""
    job_text = f"{target_title} {job_description}".casefold()
    candidate_text = _candidate_context_text(candidate_profile, candidate_answers)
    freelance_terms = (
        "freelance", "independent contractor", "consulting engagement",
        "contract consultant", "fractional", "client portfolio",
    )
    technical_terms = _DOMAIN_TERMS["technical"]

    if any(term in job_text for term in freelance_terms):
        return "freelance"
    if _candidate_freelance_score(candidate_profile, candidate_text) >= 3:
        return "freelance"
    if _looks_like_career_change(job_text, candidate_profile, candidate_text):
        return "career_changer"
    if any(term in job_text for term in technical_terms):
        return "technical"
    return "standard"


def recommend_visual_design(
    job_description: str,
    target_title: str = "",
    *,
    resume_format: str | None = None,
    career_stage: str | None = None,
    candidate_profile: Any | None = None,
    candidate_answers: Iterable[Any] | None = None,
) -> str:
    """Recommend presentation using target-industry and candidate context."""
    job_text = f"{target_title} {job_description}".casefold()
    candidate_text = _candidate_context_text(candidate_profile, candidate_answers)
    conservative_terms = (
        "bank", "banking", "financial", "audit", "auditor", "government",
        "federal", "regulatory", "compliance", "legal", "insurance",
    )
    modern_context_terms = (
        "startup", "saas", "software company", "technology company", "product-led",
        "digital product", "innovation lab", "developer platform",
    )
    if any(term in job_text for term in conservative_terms):
        return "corporate"
    if any(term in job_text for term in modern_context_terms):
        return "modern"
    if normalize_resume_format(resume_format) == "technical":
        return "modern"
    if normalize_resume_format(resume_format) == "career_changer" and _domain_scores(job_text)["technical"] >= 2:
        return "modern"
    if not job_text.strip() and _domain_scores(candidate_text)["technical"] >= 2:
        return "modern"
    if normalize_career_stage(career_stage) == "executive":
        return "corporate"
    return "corporate"

def career_stage_options() -> list[dict[str, str]]:
    return [dict(option) for option in CAREER_STAGE_OPTIONS]


def resume_format_options() -> list[dict[str, str]]:
    return [dict(option) for option in RESUME_FORMAT_OPTIONS]


def visual_design_options() -> list[dict[str, str]]:
    return [dict(option) for option in VISUAL_DESIGN_OPTIONS]


def resume_preference_label(
    career_stage: str | None,
    resume_format: str | None,
    visual_design: str | None,
) -> str:
    def option_label(options, key):
        return next(option["label"] for option in options if option["key"] == key)

    stage = normalize_career_stage(career_stage)
    format_key = normalize_resume_format(resume_format)
    design = normalize_visual_design(visual_design)
    return (
        f"{option_label(CAREER_STAGE_OPTIONS, stage)} · "
        f"{option_label(RESUME_FORMAT_OPTIONS, format_key)} · "
        f"{option_label(VISUAL_DESIGN_OPTIONS, design)}"
    )


def resume_style_options() -> list[dict[str, str]]:
    """Legacy career-stage options used by existing application records."""
    return [
        {
            "key": theme.key,
            "label": theme.label,
            "audience": theme.audience,
            "description": theme.description,
            "collection": theme.collection,
            "selector_note": theme.selector_note,
        }
        for theme in RESUME_STYLE_THEMES.values()
    ]

def _set_font(
    style,
    *,
    font_name: str,
    size: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor,
) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = color
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), font_name)


def _set_character_spacing(style, value: int | None) -> None:
    r_pr = style._element.get_or_add_rPr()
    existing = r_pr.find(qn("w:spacing"))
    if value is None:
        if existing is not None:
            r_pr.remove(existing)
        return
    if existing is None:
        existing = OxmlElement("w:spacing")
        r_pr.append(existing)
    existing.set(qn("w:val"), str(value))


def _set_small_caps(style, enabled: bool) -> None:
    r_pr = style._element.get_or_add_rPr()
    existing = r_pr.find(qn("w:smallCaps"))
    if not enabled:
        if existing is not None:
            r_pr.remove(existing)
        return
    if existing is None:
        existing = OxmlElement("w:smallCaps")
        r_pr.append(existing)
    existing.set(qn("w:val"), "1")


def _get_or_add_paragraph_style(document: Document, name: str):
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles["Normal"]
    return style


def _remove_paragraph_borders(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)


def _remove_paragraph_shading(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:shd"))
    if existing is not None:
        p_pr.remove(existing)


def _set_paragraph_shading(style, *, fill: str) -> None:
    _remove_paragraph_shading(style)
    p_pr = style._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _set_left_border(style, *, color: str, size: str = "18") -> None:
    _remove_paragraph_borders(style)
    p_pr = style._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)


def _set_bottom_border(
    style,
    *,
    color: str = "1F4E79",
    size: str = "6",
    line_style: str = "single",
    space: str = "2",
) -> None:
    _remove_paragraph_borders(style)
    p_pr = style._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), line_style)
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def _set_default_language(style, language: str = "en-US") -> None:
    r_pr = style._element.get_or_add_rPr()
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), language)
    lang.set(qn("w:eastAsia"), language)


def configure_resume_document(
    document: Document,
    style_key: str = DEFAULT_RESUME_STYLE,
    *,
    career_stage: str | None = None,
    visual_design: str | None = None,
) -> ResumeStyleTheme:
    """Configure page geometry from independent stage and design choices."""
    stage = normalize_career_stage(career_stage or style_key)
    theme = compose_resume_theme(stage, visual_design)
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(theme.top_margin)
        section.bottom_margin = Inches(theme.bottom_margin)
        section.left_margin = Inches(theme.left_margin)
        section.right_margin = Inches(theme.right_margin)
        section.header_distance = Inches(0.2)
        section.footer_distance = Inches(0.2)
        columns = section._sectPr.find(qn("w:cols"))
        if columns is None:
            columns = OxmlElement("w:cols")
            section._sectPr.append(columns)
        for child in list(columns):
            columns.remove(child)
        columns.set(qn("w:num"), "1")

    normal = document.styles["Normal"]
    _set_font(
        normal,
        font_name=theme.body_font,
        size=theme.body_size,
        color=theme.text_color,
    )
    _set_default_language(normal)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_NAME)
    _set_font(
        style,
        font_name=theme.heading_font,
        size=theme.name_size,
        bold=not theme.is_mid_career_corporate,
        color=theme.accent_color,
    )
    _set_character_spacing(style, 100 if theme.is_mid_career_corporate else None)
    _set_small_caps(style, False)
    style.paragraph_format.alignment = theme.header_alignment
    style.paragraph_format.space_after = Pt(3 if theme.is_mid_career_corporate else 1)
    style.paragraph_format.keep_with_next = True
    _remove_paragraph_borders(style)
    _remove_paragraph_shading(style)

    style = _get_or_add_paragraph_style(document, STYLE_TARGET_TITLE)
    _set_font(
        style,
        font_name=theme.heading_font,
        size=theme.target_size,
        bold=not theme.is_mid_career_corporate,
        color=theme.text_color if theme.is_mid_career_corporate else theme.accent_color,
    )
    _set_character_spacing(style, 100 if theme.is_mid_career_corporate else None)
    _set_small_caps(style, False)
    style.paragraph_format.alignment = theme.header_alignment
    style.paragraph_format.space_after = Pt(3 if theme.is_mid_career_corporate else 1)
    style.paragraph_format.keep_with_next = True
    if theme.is_mid_career_corporate:
        _set_bottom_border(
            style,
            color="000000",
            size="4",
            line_style="double",
            space="1",
        )
    else:
        _remove_paragraph_borders(style)

    style = _get_or_add_paragraph_style(document, STYLE_CONTACT)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.contact_size,
        color=theme.text_color,
    )
    _set_character_spacing(style, 26 if theme.is_mid_career_corporate else None)
    _set_small_caps(style, False)
    style.paragraph_format.alignment = theme.header_alignment
    style.paragraph_format.space_after = Pt(2 if theme.is_mid_career_corporate else 5)
    style.paragraph_format.keep_with_next = True
    _remove_paragraph_shading(style)
    if theme.is_executive and theme.is_corporate:
        _set_bottom_border(style, color="253752", size="10")
    else:
        _remove_paragraph_borders(style)

    style = _get_or_add_paragraph_style(document, STYLE_SECTION_HEADING)
    _set_font(
        style,
        font_name=theme.heading_font,
        size=theme.section_size,
        bold=True,
        color=theme.accent_color,
    )
    _set_character_spacing(style, 60 if theme.is_mid_career_corporate else None)
    _set_small_caps(style, theme.is_mid_career_corporate)
    style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if theme.is_mid_career_corporate
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    style.paragraph_format.space_before = Pt(theme.section_space_before)
    style.paragraph_format.space_after = Pt(theme.section_space_after)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.widow_control = True
    if theme.is_early_career and theme.is_corporate:
        _set_left_border(style, color="2563EB", size="14")
        _set_paragraph_shading(style, fill="EFF6FF")
        style.paragraph_format.left_indent = Inches(0.07)
        style.paragraph_format.right_indent = Inches(0.03)
    elif theme.section_border:
        accent_hex = "{:02X}{:02X}{:02X}".format(*theme.accent_color)
        _set_bottom_border(style, color=accent_hex)
        _remove_paragraph_shading(style)
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
    else:
        _remove_paragraph_borders(style)
        _remove_paragraph_shading(style)
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)

    style = _get_or_add_paragraph_style(document, STYLE_SUMMARY)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.body_size,
        color=theme.text_color,
    )
    style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
        if theme.is_mid_career_corporate
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    style.paragraph_format.space_after = Pt(2 if theme.is_mid_career_corporate else 3)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_SKILL_LINE)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.skill_size,
        color=theme.text_color,
    )
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(0 if theme.is_mid_career_corporate else 1)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_EMPLOYER_LINE)
    _set_font(
        style,
        font_name=(theme.heading_font if theme.is_executive else theme.body_font),
        size=theme.employer_size,
        color=theme.text_color,
    )
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_before = Pt(7 if theme.is_mid_career_corporate else 4.5 if theme.is_executive else 4)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_ROLE)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.role_size,
        bold=True,
        color=(theme.accent_color if (theme.is_executive or theme.is_early_career) else theme.text_color),
    )
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(0 if theme.is_mid_career_corporate else 1)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_BULLET)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.bullet_size,
        color=theme.text_color,
    )
    style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
        if theme.is_mid_career_corporate
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    style.paragraph_format.left_indent = Inches(0.22 if theme.is_mid_career_corporate else 0.18)
    style.paragraph_format.first_line_indent = Inches(-0.17 if theme.is_mid_career_corporate else -0.13)
    style.paragraph_format.space_after = Pt(theme.bullet_space_after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_EDUCATION)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.education_size,
        color=theme.text_color,
    )
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_before = Pt(1 if theme.is_mid_career_corporate else 2)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_EDUCATION_META)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.education_size - 0.25,
        color=theme.text_color,
    )
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.widow_control = True

    style = _get_or_add_paragraph_style(document, STYLE_EDUCATION_DETAIL)
    _set_font(
        style,
        font_name=theme.body_font,
        size=theme.education_size - 0.25,
        italic=not theme.is_mid_career_corporate,
        color=theme.text_color,
    )
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.left_indent = Inches(0.22 if theme.is_mid_career_corporate else 0.18)
    style.paragraph_format.first_line_indent = Inches(-0.17 if theme.is_mid_career_corporate else 0)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.widow_control = True
    return theme


def clear_document_body(document: Document) -> None:
    """Remove all body content while retaining section properties and style definitions."""
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def clear_headers_and_footers(document: Document) -> None:
    for section in document.sections:
        for part in (section.header, section.footer):
            element = part._element
            for child in list(element):
                element.remove(child)
            part.add_paragraph()
