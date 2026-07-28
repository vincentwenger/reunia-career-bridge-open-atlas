from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from resume_tailor.docx_export import TemplateError, export_resume_docx
from resume_tailor.docx_styles import (
    RESUME_STYLE_NAMES,
    STYLE_BULLET,
    STYLE_EDUCATION,
    STYLE_EMPLOYER_LINE,
    STYLE_NAME,
    STYLE_SECTION_HEADING,
)
from resume_tailor.validation import build_approved_resume


def _body_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_clean_template_contains_styles_without_candidate_content(project_root, profile):
    template_path = project_root / "data" / "resume_template_professional.docx"
    document = Document(template_path)
    text = _body_text(document)

    assert profile.name not in text
    assert profile.contact.email not in text
    assert "{{" not in text
    assert not document.tables
    for style_name in RESUME_STYLE_NAMES:
        assert style_name in document.styles

    with ZipFile(template_path) as archive:
        relationships = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".rels")
        )
    assert "linkedin.com/in/vincentwenger" not in relationships
    assert "github.com/vincentwenger" not in relationships


def test_dynamic_docx_export_builds_complete_resume(project_root, profile, analysis, proposal):
    approved = build_approved_resume(profile, analysis, proposal)
    result = export_resume_docx(project_root / "data" / "resume_template_professional.docx", profile, approved)

    document = Document(BytesIO(result))
    body_text = _body_text(document)

    assert "{{" not in body_text
    assert profile.name in body_text
    assert profile.contact.location in body_text
    assert profile.contact.phone in body_text
    assert profile.contact.email in body_text
    assert "Axiom Developer" in body_text
    assert proposal.professional_summary in body_text
    assert "Hard Skills:" in body_text
    assert "Led the end-to-end implementation" in body_text
    assert "Delivered 10 complex software projects" not in body_text
    for experience in profile.experiences:
        assert experience.employer in body_text
        assert experience.title in body_text
        assert experience.dates in body_text
    for education in profile.education:
        assert education.credential in body_text
        assert education.date in body_text

    assert document.paragraphs[0].style.name == STYLE_NAME
    assert any(paragraph.style.name == STYLE_SECTION_HEADING for paragraph in document.paragraphs)
    assert any(paragraph.style.name == STYLE_EMPLOYER_LINE for paragraph in document.paragraphs)
    assert any(paragraph.style.name == STYLE_EDUCATION for paragraph in document.paragraphs)

    bullets = [paragraph for paragraph in document.paragraphs if paragraph.style.name == STYLE_BULLET]
    assert bullets
    assert all(paragraph.text.startswith("• ") for paragraph in bullets)
    assert all(paragraph.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT) for paragraph in bullets)
    assert document.styles[STYLE_BULLET].paragraph_format.first_line_indent.inches < 0
    assert not any(paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY for paragraph in document.paragraphs)
    assert not document.tables
    assert not document.inline_shapes


def test_export_uses_profile_content_instead_of_template_hardcoding(
    project_root, profile, analysis, proposal
):
    alternate_profile = profile.model_copy(deep=True)
    alternate_profile.name = "Jordan Lee"
    alternate_profile.contact.location = "Denver, CO"
    alternate_profile.contact.phone = "303-555-0198"
    alternate_profile.contact.email = "jordan@example.com"
    alternate_profile.contact.linkedin_label = "LinkedIn"
    alternate_profile.contact.linkedin_url = "https://www.linkedin.com/in/jordan-example/"
    alternate_profile.contact.github_label = "Portfolio"
    alternate_profile.contact.github_url = "https://example.com/jordan"
    alternate_profile.experiences[0].employer = "Northstar Labs"
    alternate_profile.education[0].institution = "Example University"

    approved = build_approved_resume(alternate_profile, analysis, proposal)
    result = export_resume_docx(
        project_root / "data" / "resume_template_professional.docx", alternate_profile, approved
    )
    document = Document(BytesIO(result))
    body_text = _body_text(document)

    assert "Jordan Lee" in body_text
    assert "Denver, CO" in body_text
    assert "Northstar Labs" in body_text
    assert "Example University" in body_text
    assert "Vincent Wenger" not in body_text
    assert "Nasdaq" not in body_text

    with ZipFile(BytesIO(result)) as archive:
        relationships = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".rels")
        )
    assert "linkedin.com/in/jordan-example" in relationships
    assert "example.com/jordan" in relationships
    assert "mailto:jordan@example.com" in relationships


def test_long_first_experience_uses_dynamic_page_balance(project_root, profile, analysis, proposal):
    approved = build_approved_resume(profile, analysis, proposal)
    result = export_resume_docx(project_root / "data" / "resume_template_professional.docx", profile, approved)
    document = Document(BytesIO(result))
    employer_lines = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == STYLE_EMPLOYER_LINE
    ]

    assert len(employer_lines) >= 2
    assert employer_lines[0].paragraph_format.page_break_before in (None, False)
    assert employer_lines[1].paragraph_format.page_break_before is True


def test_docx_export_blocks_adjacent_repeated_words(
    project_root, profile, analysis, proposal
):
    broken = proposal.model_copy(deep=True)
    broken.professional_summary = broken.professional_summary.replace(
        "and client training", "and and client training"
    )
    approved = build_approved_resume(profile, analysis, broken)

    with pytest.raises(TemplateError, match="adjacent repeated word"):
        export_resume_docx(
            project_root / "data" / "resume_template_professional.docx", profile, approved
        )
