from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from resume_tailor.docx_export import export_resume_docx
from resume_tailor.models import CandidateAnswer, VerifiedSkills
from resume_tailor.resume_report import build_resume_report
from resume_tailor.docx_styles import (
    RESUME_STYLE_NAMES,
    STYLE_EMPLOYER_LINE,
    STYLE_NAME,
    STYLE_SECTION_HEADING,
    STYLE_SUMMARY,
    STYLE_TARGET_TITLE,
    career_stage_options,
    normalize_resume_style,
    recommend_career_stage,
    recommend_resume_format,
    recommend_resume_style,
    recommend_visual_design,
    resume_format_options,
    resume_style_options,
    visual_design_options,
)
from resume_tailor.validation import build_approved_resume
from resume_tailor.web_state import WorkflowState


STYLE_KEYS = ("early_career", "professional", "executive")


def _body_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _section_index(document: Document, heading: str) -> int:
    normalized = heading.casefold()
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().casefold() == normalized:
            return index
    raise AssertionError(f"Section not found: {heading}")


def test_all_style_templates_are_clean_and_share_named_styles(project_root, profile):
    for style in STYLE_KEYS:
        template_path = project_root / "data" / f"resume_template_{style}.docx"
        assert template_path.exists()
        document = Document(template_path)
        text = _body_text(document)
        assert profile.name not in text
        assert profile.contact.email not in text
        assert not document.tables
        for style_name in RESUME_STYLE_NAMES:
            assert style_name in document.styles


def test_only_three_template_files_remain(project_root):
    template_names = sorted(
        path.name for path in (project_root / "data").glob("resume_template_*.docx")
    )
    assert template_names == sorted(
        [
            "resume_template_early_career.docx",
            "resume_template_professional.docx",
            "resume_template_executive.docx",
        ]
    )


def test_all_exports_keep_content_but_change_presentation(
    project_root, profile, analysis, proposal
):
    approved = build_approved_resume(profile, analysis, proposal)
    exports = {
        style: export_resume_docx(
            project_root / "data" / f"resume_template_{style}.docx",
            profile,
            approved,
            style_key=style,
        )
        for style in STYLE_KEYS
    }

    assert len(set(exports.values())) == len(STYLE_KEYS)
    documents = {style: Document(BytesIO(payload)) for style, payload in exports.items()}

    for expected in (
        profile.name,
        proposal.professional_summary,
        "Led the end-to-end implementation",
        "Axiom regulatory reporting platform",
    ):
        for document in documents.values():
            assert expected in _body_text(document)

    assert documents["early_career"].styles[STYLE_NAME].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert documents["professional"].styles[STYLE_NAME].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert documents["executive"].styles[STYLE_NAME].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT

    headings = {
        style: next(
            p for p in document.paragraphs if p.style.name == STYLE_SECTION_HEADING
        )
        for style, document in documents.items()
    }
    assert {paragraph.text for paragraph in headings.values()} == {"Professional Summary"}

    early_properties = documents["early_career"].styles[STYLE_SECTION_HEADING]._element.get_or_add_pPr()
    early_border = early_properties.find(qn("w:pBdr"))
    early_shading = early_properties.find(qn("w:shd"))
    assert early_border is not None
    assert early_border.find(qn("w:left")) is not None
    assert early_shading is not None
    assert early_shading.get(qn("w:fill")) == "EFF6FF"

    professional_properties = documents["professional"].styles[STYLE_SECTION_HEADING]._element.get_or_add_pPr()
    assert professional_properties.find(qn("w:pBdr")) is None
    assert documents["professional"].styles[STYLE_SECTION_HEADING].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    professional_heading_rpr = documents["professional"].styles[STYLE_SECTION_HEADING]._element.get_or_add_rPr()
    assert professional_heading_rpr.find(qn("w:smallCaps")) is not None
    assert professional_heading_rpr.find(qn("w:spacing")).get(qn("w:val")) == "60"

    executive_properties = documents["executive"].styles[STYLE_SECTION_HEADING]._element.get_or_add_pPr()
    executive_border = executive_properties.find(qn("w:pBdr"))
    assert executive_border is not None
    assert executive_border.find(qn("w:bottom")) is not None

    assert _section_index(documents["early_career"], "Education and Professional Development") < _section_index(
        documents["early_career"], "Work Experience"
    )
    for style in ("professional", "executive"):
        assert _section_index(documents[style], "Work Experience") < _section_index(
            documents[style], "Education and Professional Development"
        )

    expected_subjects = {
        "early_career": "Early Career · Standard Professional · Corporate",
        "professional": "Mid-Career Professional · Standard Professional · Corporate",
        "executive": "Executive Leadership · Standard Professional · Corporate",
    }
    for style, phrase in expected_subjects.items():
        assert phrase in documents[style].core_properties.subject


def test_mid_career_template_recreates_original_resume_style(
    project_root, profile, analysis, proposal
):
    approved = build_approved_resume(profile, analysis, proposal)
    payload = export_resume_docx(
        project_root / "data" / "resume_template_professional.docx",
        profile,
        approved,
        style_key="professional",
    )
    document = Document(BytesIO(payload))

    name_style = document.styles[STYLE_NAME]
    assert name_style.font.name == "Calibri"
    assert name_style.font.size.pt == 15
    assert name_style.font.bold is False
    name_spacing = name_style._element.get_or_add_rPr().find(qn("w:spacing"))
    assert name_spacing is not None
    assert name_spacing.get(qn("w:val")) == "100"

    title_border = document.styles[STYLE_TARGET_TITLE]._element.get_or_add_pPr().find(qn("w:pBdr"))
    assert title_border is not None
    assert title_border.find(qn("w:bottom")).get(qn("w:val")) == "double"
    assert document.styles[STYLE_SUMMARY].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    nonempty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    assert nonempty[0].text == profile.name
    assert nonempty[1].text == approved.target_title
    assert nonempty[2].text == profile.contact.location
    assert profile.contact.phone in nonempty[3].text
    assert profile.contact.email in nonempty[3].text

    employer_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == STYLE_EMPLOYER_LINE
    )
    assert employer_paragraph.runs[0].text == profile.experiences[0].employer
    assert employer_paragraph.runs[0].underline is True
    assert employer_paragraph.runs[0].bold is not True

    first_education = profile.education[0]
    education_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if first_education.credential in paragraph.text
    )
    assert first_education.institution in education_paragraph.text
    assert first_education.location in education_paragraph.text


def test_style_recommendation_helper():
    assert recommend_resume_style(
        "Build Python APIs and cloud software for a SaaS product", "Software Engineer"
    ) == "professional"
    assert recommend_resume_style(
        "Support regulatory compliance, risk controls, and banking audits", "IT Auditor"
    ) == "professional"
    assert recommend_resume_style(
        "Own enterprise strategy, portfolio direction, and organizational transformation",
        "Vice President of Technology",
    ) == "executive"
    assert recommend_resume_style(
        "Join our graduate program and learn from experienced engineers",
        "Entry-Level Software Engineer",
    ) == "early_career"


def test_unknown_style_uses_mid_career_default():
    assert normalize_resume_style("removed-template") == "professional"


def test_style_options_expose_exactly_three_choices():
    options = resume_style_options()
    assert [option["key"] for option in options] == list(STYLE_KEYS)
    assert [option["label"] for option in options] == [
        "Early Career",
        "Mid-Career Professional",
        "Executive Leadership",
    ]
    assert [option["collection"] for option in options] == ["career_stage"] * 3
    assert "original resume" in options[1]["selector_note"].casefold()


def test_workflow_state_defaults_to_composable_resume_preferences(profile):
    state = WorkflowState(source_profile=profile)
    assert state.resume_style == "professional"
    assert state.resume_style_explicit is False
    assert state.resume_career_stage == "mid_career"
    assert state.resume_format == "standard"
    assert state.resume_visual_design == "corporate"


def test_composable_option_sets_expose_recommended_choices():
    assert [option["key"] for option in career_stage_options()] == [
        "early_career", "mid_career", "executive"
    ]
    assert [option["key"] for option in resume_format_options()] == [
        "standard", "technical", "career_changer", "freelance"
    ]
    assert [option["key"] for option in visual_design_options()] == [
        "corporate", "modern"
    ]


def test_format_and_design_recommendations_are_independent():
    assert recommend_career_stage(
        "Build Python APIs and cloud software", "Senior Software Engineer"
    ) == "mid_career"
    assert recommend_resume_format(
        "Build Python APIs and cloud software", "Senior Software Engineer"
    ) == "technical"
    assert recommend_visual_design(
        "Build Python APIs and cloud software",
        "Senior Software Engineer",
        resume_format="technical",
    ) == "modern"
    assert recommend_visual_design(
        "Build regulatory reports for a bank",
        "Axiom Developer",
        resume_format="technical",
    ) == "corporate"


def test_candidate_experience_refines_career_stage_recommendation(profile):
    assert recommend_career_stage(
        "Entry-level role for recent graduates",
        "Junior Software Engineer",
    ) == "early_career"
    assert recommend_career_stage(
        "Entry-level role for recent graduates",
        "Junior Software Engineer",
        candidate_profile=profile,
    ) == "mid_career"
    assert recommend_career_stage(
        "Own enterprise strategy and organizational transformation",
        "Vice President of Technology",
        candidate_profile=profile,
    ) == "executive"


def test_candidate_background_can_trigger_career_changer_format(profile):
    assert recommend_resume_format(
        "Lead brand campaigns, demand generation, and customer acquisition",
        "Marketing Manager",
        candidate_profile=profile,
    ) == "career_changer"
    assert recommend_resume_format(
        "Build Python APIs, cloud services, and SQL databases",
        "Software Engineer",
        candidate_profile=profile,
    ) == "technical"


def test_confirmed_experience_can_replace_career_changer_with_technical(profile):
    accounting_experience = profile.experiences[0].model_copy(
        update={
            "title": "Senior Accountant",
            "employer": "Example Finance",
            "dates": "01/2015 - 12/2024",
            "bullets": [
                profile.experiences[0].bullets[0].model_copy(
                    update={
                        "text": "Prepared financial statements and account reconciliations."
                    }
                )
            ],
        }
    )
    accounting_profile = profile.model_copy(
        update={
            "current_summary": "Accountant with 10 years of financial reporting experience.",
            "skills": VerifiedSkills(),
            "education": [],
            "experiences": [accounting_experience],
        }
    )
    job_description = "Build Python APIs on AWS and design SQL databases."
    assert recommend_resume_format(
        job_description,
        "Software Engineer",
        candidate_profile=accounting_profile,
    ) == "career_changer"

    answers = [
        CandidateAnswer(
            question_id="Q1",
            question="Do you have hands-on software engineering experience?",
            answer_type="yes_no_with_details",
            yes_no=True,
            text="Built Python APIs on AWS and designed SQL databases.",
        )
    ]
    assert recommend_resume_format(
        job_description,
        "Software Engineer",
        candidate_profile=accounting_profile,
        candidate_answers=answers,
    ) == "technical"


def test_multiple_independent_engagements_trigger_freelance_format(profile):
    engagements = [
        experience.model_copy(
            update={
                "title": "Independent Consultant",
                "employer": f"Client {index}",
            }
        )
        for index, experience in enumerate(profile.experiences[:2], start=1)
    ]
    consulting_profile = profile.model_copy(update={"experiences": engagements})
    assert recommend_resume_format(
        "Deliver cross-functional transformation projects",
        "Program Manager",
        candidate_profile=consulting_profile,
    ) == "freelance"


def test_resume_formats_change_section_structure_without_changing_content(
    project_root, profile, analysis, proposal
):
    approved = build_approved_resume(profile, analysis, proposal)
    template = project_root / "data" / "resume_template_professional.docx"
    payloads = {
        format_key: export_resume_docx(
            template,
            profile,
            approved,
            career_stage="mid_career",
            resume_format=format_key,
            visual_design="corporate",
        )
        for format_key in ("standard", "technical", "career_changer", "freelance")
    }
    documents = {key: Document(BytesIO(value)) for key, value in payloads.items()}
    for document in documents.values():
        assert proposal.professional_summary in _body_text(document)
        assert "Led the end-to-end implementation" in _body_text(document)

    technical = documents["technical"]
    assert _section_index(technical, "Technical Skills") < _section_index(
        technical, "Technical Profile"
    )
    assert _section_index(technical, "Engineering Experience") < _section_index(
        technical, "Education and Certifications"
    )

    career_changer = documents["career_changer"]
    assert _section_index(
        career_changer, "Transferable and Relevant Skills"
    ) < _section_index(career_changer, "Relevant Experience")
    assert _section_index(
        career_changer, "Education and Professional Development"
    ) < _section_index(career_changer, "Relevant Experience")

    freelance = documents["freelance"]
    assert _section_index(freelance, "Client and Project Experience") > _section_index(
        freelance, "Core Capabilities"
    )


def test_visual_design_changes_presentation_independently(
    project_root, profile, analysis, proposal
):
    approved = build_approved_resume(profile, analysis, proposal)
    template = project_root / "data" / "resume_template_professional.docx"
    corporate = Document(BytesIO(export_resume_docx(
        template, profile, approved, career_stage="mid_career",
        resume_format="technical", visual_design="corporate"
    )))
    modern = Document(BytesIO(export_resume_docx(
        template, profile, approved, career_stage="mid_career",
        resume_format="technical", visual_design="modern"
    )))
    corporate_headings = [
        paragraph.text for paragraph in corporate.paragraphs
        if paragraph.style.name == STYLE_SECTION_HEADING
    ]
    modern_headings = [
        paragraph.text for paragraph in modern.paragraphs
        if paragraph.style.name == STYLE_SECTION_HEADING
    ]
    assert corporate_headings == modern_headings
    for expected in (profile.name, proposal.professional_summary, "Led the end-to-end implementation"):
        assert expected in _body_text(corporate)
        assert expected in _body_text(modern)
    assert corporate.styles[STYLE_NAME].font.name == "Calibri"
    assert corporate.styles[STYLE_NAME].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert modern.styles[STYLE_NAME].font.name == "Arial"
    assert modern.styles[STYLE_NAME].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT




def test_all_composable_combinations_export_and_report_valid_structure(
    project_root, profile, analysis, proposal
):
    for stage, template_style in (
        ("early_career", "early_career"),
        ("mid_career", "professional"),
        ("executive", "executive"),
    ):
        template = project_root / "data" / f"resume_template_{template_style}.docx"
        for format_key in ("standard", "technical", "career_changer", "freelance"):
            for design in ("corporate", "modern"):
                approved = build_approved_resume(profile, analysis, proposal)
                payload = export_resume_docx(
                    template,
                    profile,
                    approved,
                    career_stage=stage,
                    resume_format=format_key,
                    visual_design=design,
                )
                assert payload.startswith(b"PK")
                report = build_resume_report(
                    profile,
                    analysis,
                    proposal,
                    generated_filename="combination.docx",
                    template_path=template,
                    resume_title=analysis.target_title,
                    generated_document_bytes=payload,
                    career_stage=stage,
                    resume_format=format_key,
                    visual_design=design,
                )
                structure_check = next(
                    check
                    for section in report.sections()
                    if section.name == "Content Quality"
                    for subsection in section.subsections
                    if subsection.name == "Data & Structure"
                    for check in subsection.checks
                    if check.label == "The generated resume preserves the expected section hierarchy"
                )
                assert structure_check.status == "pass", (stage, format_key, design, structure_check.detail)


def test_step_four_exposes_composable_resume_preferences(project_root: Path):
    template = (project_root / "templates" / "index.html").read_text(encoding="utf-8")
    app_source = (project_root / "app.py").read_text(encoding="utf-8")

    assert 'id="resume-style-selector"' in template
    assert "Career stage" in template
    assert "Resume format" in template
    assert "Visual design" in template
    assert "Technical / Engineering" not in template  # labels come from backend options
    assert 'data-auto-submit-radio="career_stage"' in template
    assert 'data-auto-submit-radio="resume_format"' in template
    assert 'data-auto-submit-radio="visual_design"' in template
    assert "Apply selected style" not in template
    assert "Selecting any option applies it immediately" in template
    assert "AI optimization are not rerun" in template
    assert "uploaded resume, and any experience you confirmed" in template

    javascript = (project_root / "static" / "app.js").read_text(encoding="utf-8")
    assert "form[data-auto-submit-radio]" in javascript
    assert "form.requestSubmit()" in javascript
    assert '@app.post("/resume-style")' in app_source
    for style in STYLE_KEYS:
        assert f"resume_template_{style}.docx" in app_source
    assert "resume_format_options" in app_source
    assert "visual_design_options" in app_source
    assert "Resume Reports were refreshed without rerunning optimization" in app_source

