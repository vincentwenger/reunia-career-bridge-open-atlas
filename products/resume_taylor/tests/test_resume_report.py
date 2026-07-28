from __future__ import annotations

import base64

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from resume_tailor.models import CandidateAnswer, EvidenceMatch, JobAnalysis, JobRequirement
from resume_tailor.resume_report import (
    build_evidence_gap_report,
    build_initial_resume_proposal,
    build_resume_report,
    initial_resume_title,
)


def _all_checks(section):
    return [check for subsection in section.subsections for check in subsection.checks]


def test_resume_report_contains_requested_sections(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )

    assert [section.name for section in report.sections()] == [
        "Hard skills",
        "Evidence & Gaps",
        "Content Quality",
        "Searchability",
        "Recruiter tips",
        "Formatting",
        "Soft skills",
    ]
    assert [subsection.name for subsection in report.searchability.subsections] == [
        "Contact Information",
        "Summary",
        "Section Headings",
        "Job Title Match",
        "Date Formatting",
        "Education Match",
        "File Type",
    ]


def test_resume_report_sections_identify_primary_workflow_ownership(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    ownership = {
        section.name: (
            section.primary_workflow_step,
            section.primary_workflow_stage,
            section.verification_workflow_step,
            section.verification_workflow_stage,
        )
        for section in report.sections()
    }
    assert ownership == {
        "Hard skills": ("Review Job Alignment", "draft", "", ""),
        "Evidence & Gaps": ("Review Job Alignment", "draft", "", ""),
        "Content Quality": ("Optimize & Export", "final", "", ""),
        "Searchability": ("Optimize & Export", "final", "", ""),
        "Recruiter tips": ("Optimize & Export", "final", "", ""),
        "Formatting": ("Optimize & Export", "final", "", ""),
        "Soft skills": ("Optimize & Export", "final", "", ""),
    }

def test_searchability_passes_core_profile_checks(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    checks = _all_checks(report.searchability)
    by_label = {check.label: check for check in checks}

    assert by_label["You provided your email"].status == "pass"
    assert by_label["You provided your phone number"].status == "pass"
    assert by_label['We found an "Education" section in your resume'].status == "pass"
    assert by_label["Work-experience dates are properly formatted"].status == "pass"
    assert by_label["You are using a .docx resume"].status == "warning"


def test_unsupported_critical_hard_skill_is_flagged(project_root, profile, analysis, proposal):
    analysis.requirements.append(
        JobRequirement(
            id="R4",
            category="technical_skill",
            priority="critical",
            requirement="Build COBOL batch applications",
            keywords=["COBOL"],
        )
    )
    proposal.evidence_matches.append(
        EvidenceMatch(
            requirement_id="R4",
            status="unsupported",
            evidence_ids=[],
            rationale="No COBOL evidence in the source profile.",
        )
    )

    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    hard_checks = _all_checks(report.hard_skills)
    cobol_check = next(check for check in hard_checks if "COBOL" in check.label)
    recruiter_checks = _all_checks(report.recruiter_tips)
    critical_gap_check = next(
        check for check in recruiter_checks if check.label == "Critical gaps are disclosed instead of invented"
    )

    assert cobol_check.status == "fail"
    assert critical_gap_check.status == "warning"
    assert "COBOL" in critical_gap_check.detail


def test_invalid_work_date_is_a_searchability_failure(project_root, profile, analysis, proposal):
    changed_profile = profile.model_copy(deep=True)
    changed_profile.experiences[0].dates = "September 2013 through May 2025"

    report = build_resume_report(
        changed_profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    date_check = next(
        check
        for check in _all_checks(report.searchability)
        if check.label == "Work-experience dates are properly formatted"
    )

    assert date_check.status == "fail"


def test_recruiter_tips_contains_requested_subsections(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )

    assert [subsection.name for subsection in report.recruiter_tips.subsections] == [
        "Job Level Match",
        "Measurable Results",
        "Resume Tone",
        "Web Presence",
        "Word Count",
    ]


def test_recruiter_tips_passes_requested_checks(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    by_label = {check.label: check for check in _all_checks(report.recruiter_tips)}

    assert by_label["Your years of experience align with the role's requirements"].status == "pass"
    assert by_label["There are five or more mentions of measurable results"].status == "pass"
    assert by_label["The resume tone is positive and avoids common clichés and buzzwords"].status == "pass"
    assert by_label["You linked to a website that builds your web credibility"].status == "pass"
    assert by_label["The resume contains fewer than 1,000 words"].status == "pass"


def test_job_level_match_flags_insufficient_years(project_root, profile, analysis, proposal):
    analysis.requirements.append(
        JobRequirement(
            id="R5",
            category="qualification",
            priority="critical",
            requirement="Requires at least 20 years of professional experience",
            keywords=["20 years"],
        )
    )

    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    check = next(
        item
        for item in _all_checks(report.recruiter_tips)
        if item.label == "Your years of experience align with the role's requirements"
    )

    assert check.status == "fail"
    assert "20 years" in check.detail


def test_resume_tone_flags_common_cliche(project_root, profile, analysis, proposal):
    changed_proposal = proposal.model_copy(deep=True)
    changed_proposal.professional_summary = (
        "Results-driven team player with a proven track record of delivering software solutions."
    )

    report = build_resume_report(
        profile,
        analysis,
        changed_proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    check = next(
        item
        for item in _all_checks(report.recruiter_tips)
        if item.label == "The resume tone is positive and avoids common clichés and buzzwords"
    )

    assert check.status == "warning"
    assert "results driven" in check.detail
    assert "team player" in check.detail


def test_word_count_fails_at_one_thousand_words(project_root, profile, analysis, proposal):
    changed_proposal = proposal.model_copy(deep=True)
    changed_proposal.professional_summary = " ".join(["engineer"] * 1000)

    report = build_resume_report(
        profile,
        analysis,
        changed_proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    check = next(
        item
        for item in _all_checks(report.recruiter_tips)
        if item.label == "The resume contains fewer than 1,000 words"
    )

    assert check.status == "fail"


def test_formatting_contains_requested_subsections_and_checks(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )

    assert [subsection.name for subsection in report.formatting.subsections] == [
        "Layout",
        "Font Check",
        "Typography Consistency",
        "Page Setup",
    ]

    labels = {check.label for check in _all_checks(report.formatting)}
    assert {
        "Your resume does not contain columns",
        "Your paragraphs are not longer than 40 words",
        "Your resume does not contain images",
        "Your resume does not contain any tables",
        "Your resume primarily uses standardized left alignment for text sections",
        "Special characters were not overused in your resume",
        "Your resume does not contain too much bold styling",
        "All parts of the resume use an easy-to-read font color",
        "Your resume does not overuse different fonts",
        "Your resume does not contain non-standard fonts",
        "The average font size meets readability and ATS standards",
        "Your resume does not contain information in footers",
        "Your resume does not contain information in headers",
        "Your margin sizes are consistent and use standard dimensions",
        "Your document page size is standard",
        "The resume fits within the 2-page limit",
        "Bullet and list styles are consistent",
        "Bullet indentation is consistent",
        "Line spacing is consistent throughout the resume",
        "Paragraph spacing is consistent throughout the resume",
        "Section-heading typography is consistent",
    }.issubset(labels)


def test_formatting_detects_ats_unfriendly_word_features(
    tmp_path,
    project_root,
    profile,
    analysis,
    proposal,
):
    document = Document(project_root / "data" / "resume_template_professional.docx")
    document.add_table(rows=1, cols=2)

    tiny_png = tmp_path / "tiny.png"
    tiny_png.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZQmcAAAAASUVORK5CYII="
        )
    )
    document.add_paragraph().add_run().add_picture(str(tiny_png), width=Inches(0.1))

    columns = document.sections[0]._sectPr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        document.sections[0]._sectPr.append(columns)
    columns.set(qn("w:num"), "2")

    document.sections[0].header.paragraphs[0].text = "Candidate contact information"
    document.sections[0].footer.paragraphs[0].text = "Confidential resume"
    document.sections[0].left_margin = Inches(0.2)
    document.sections[0].page_width = Inches(7.0)
    document.sections[0].page_height = Inches(10.0)
    document.paragraphs[0].runs[0].font.name = "Comic Sans MS"

    changed_template = tmp_path / "ats_unfriendly_template.docx"
    document.save(changed_template)

    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=changed_template,
    )
    by_label = {check.label: check for check in _all_checks(report.formatting)}

    # The exporter now treats the template as a style pack and rebuilds the body,
    # page geometry, headers, and columns. Unsafe template content cannot leak into
    # the generated resume.
    assert by_label["Your resume does not contain columns"].status == "pass"
    assert by_label["Your resume does not contain images"].status == "pass"
    assert by_label["Your resume does not contain any tables"].status == "pass"
    assert by_label["Your resume does not contain information in headers"].status == "pass"
    assert by_label["Your resume does not contain information in footers"].status == "pass"
    assert by_label["Your margin sizes are consistent and use standard dimensions"].status == "pass"
    assert by_label["Your document page size is standard"].status == "pass"
    assert by_label["Your resume does not contain non-standard fonts"].status == "pass"


def test_hard_skills_report_compares_exact_counts_and_frequency_order(
    project_root,
    profile,
    analysis,
    proposal,
):
    job_description = (
        "Axiom regulatory reporting requires SQL and data transformation. "
        "SQL is used again for performance work. Regulatory reporting is central. "
        "Experience with COBOL is preferred."
    )
    analysis.requirements.append(
        JobRequirement(
            id="R4",
            category="technical_skill",
            priority="secondary",
            requirement="Use COBOL for legacy systems",
            keywords=["COBOL"],
        )
    )
    proposal.evidence_matches.append(
        EvidenceMatch(
            requirement_id="R4",
            status="unsupported",
            evidence_ids=[],
            rationale="No COBOL evidence.",
        )
    )

    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        job_description=job_description,
    )

    assert "Hard skills enable you" in report.hard_skills.intro
    assert "exact spelling" in report.hard_skills.intro
    comparison = report.hard_skills.subsections[0]
    assert comparison.name == "Skill comparison"

    labels = [check.label for check in comparison.checks]
    assert labels.index("SQL") < labels.index("COBOL")

    by_label = {check.label: check for check in comparison.checks}
    assert "Resume:" in by_label["SQL"].detail
    assert "Job description: 2" in by_label["SQL"].detail
    assert by_label["SQL"].status == "pass"
    assert "Resume: 0, Job description: 1" in by_label["COBOL"].detail
    assert by_label["COBOL"].status == "fail"


def test_hard_skill_count_requires_exact_spelling(
    project_root,
    profile,
    analysis,
    proposal,
):
    analysis.requirements.append(
        JobRequirement(
            id="R4",
            category="technical_skill",
            priority="important",
            requirement="Create data-driven insights",
            keywords=["data-driven insights"],
        )
    )
    proposal.evidence_matches.append(
        EvidenceMatch(
            requirement_id="R4",
            status="supported",
            evidence_ids=["NAS-01"],
            rationale="Supported for test purposes.",
        )
    )
    changed_proposal = proposal.model_copy(deep=True)
    changed_proposal.professional_summary += " Creates data driven insights for stakeholders."

    report = build_resume_report(
        profile,
        analysis,
        changed_proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        job_description="The role creates data-driven insights from reporting data.",
    )
    check = next(
        item
        for item in _all_checks(report.hard_skills)
        if item.label == "data-driven insights"
    )

    assert check.status == "fail"
    assert "Resume: 0, Job description: 1" in check.detail


def test_soft_skills_report_compares_exact_counts_and_frequency_order(
    project_root,
    profile,
    analysis,
    proposal,
):
    analysis.requirements.extend(
        [
            JobRequirement(
                id="R4",
                category="responsibility",
                priority="important",
                requirement="Use communication skills with stakeholders",
                keywords=["communication skills"],
            ),
            JobRequirement(
                id="R5",
                category="leadership",
                priority="important",
                requirement="Demonstrate executive leadership",
                keywords=["executive leadership"],
            ),
            JobRequirement(
                id="R6",
                category="qualification",
                priority="critical",
                requirement="Apply analytical judgment",
                keywords=["analytical"],
            ),
            JobRequirement(
                id="R7",
                category="leadership",
                priority="secondary",
                requirement="Coach team members",
                keywords=["coach"],
            ),
        ]
    )
    proposal.evidence_matches.extend(
        [
            EvidenceMatch(
                requirement_id="R4",
                status="supported",
                evidence_ids=["NAS-01"],
                rationale="The source profile documents stakeholder communication.",
            ),
            EvidenceMatch(
                requirement_id="R5",
                status="supported",
                evidence_ids=["NAS-02"],
                rationale="Supported for test purposes.",
            ),
            EvidenceMatch(
                requirement_id="R6",
                status="supported",
                evidence_ids=["AVI-01"],
                rationale="The source profile documents analytical work.",
            ),
            EvidenceMatch(
                requirement_id="R7",
                status="unsupported",
                evidence_ids=[],
                rationale="No verified coaching evidence.",
            ),
        ]
    )
    changed_proposal = proposal.model_copy(deep=True)
    changed_proposal.professional_summary += (
        " Analytical analytical analytical. Communication skills and communication skills. "
        "Executive leadership."
    )
    job_description = (
        "Analytical judgment is essential. Analytical thinking supports decisions, and analytical reviews are frequent. "
        "Communication skills are required. Executive leadership is valued. The manager will coach colleagues."
    )

    report = build_resume_report(
        profile,
        analysis,
        changed_proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        job_description=job_description,
    )

    assert "Soft skills are your traits" in report.soft_skills.intro
    assert "medium impact" in report.soft_skills.intro
    assert "Prioritize hard skills" in report.soft_skills.intro
    comparison = report.soft_skills.subsections[0]
    assert comparison.name == "Skill comparison"

    labels = [check.label for check in comparison.checks]
    assert labels.index("Analytical") < labels.index("coach")

    by_label = {check.label.casefold(): check for check in comparison.checks}
    assert by_label["analytical"].status == "pass"
    assert "Resume: 3, Job description: 3" in by_label["analytical"].detail
    assert by_label["communication skills"].status == "pass"
    assert "Resume: 2, Job description: 1" in by_label["communication skills"].detail
    assert by_label["executive leadership"].status == "pass"
    assert "Resume: 2, Job description: 1" in by_label["executive leadership"].detail
    assert by_label["coach"].status == "fail"
    assert "Resume: 0, Job description: 1" in by_label["coach"].detail


def test_soft_skill_count_requires_exact_spelling(
    project_root,
    profile,
    analysis,
    proposal,
):
    analysis.requirements.append(
        JobRequirement(
            id="R4",
            category="responsibility",
            priority="important",
            requirement="Demonstrate time-management",
            keywords=["time-management"],
        )
    )
    proposal.evidence_matches.append(
        EvidenceMatch(
            requirement_id="R4",
            status="supported",
            evidence_ids=["NAS-01"],
            rationale="Supported for test purposes.",
        )
    )
    changed_proposal = proposal.model_copy(deep=True)
    changed_proposal.professional_summary += " Demonstrates time management across competing priorities."

    report = build_resume_report(
        profile,
        analysis,
        changed_proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        job_description="Strong time-management is required for this role.",
    )
    check = next(
        item
        for item in _all_checks(report.soft_skills)
        if item.label == "time-management"
    )

    assert check.status == "fail"
    assert "Resume: 0, Job description: 1" in check.detail

def test_initial_resume_proposal_preserves_candidate_profile(profile):
    initial = build_initial_resume_proposal(profile)

    assert initial.professional_summary == profile.current_summary
    assert initial.skills.hard_skills == profile.skills.hard_skills
    assert initial.skills.soft_skills == profile.skills.soft_skills
    assert initial.skills.tools_software == profile.skills.tools_software
    assert initial.skills.industry_knowledge == profile.skills.industry_knowledge
    assert len(initial.bullet_proposals) == len(profile.bullet_lookup())
    assert all(item.include for item in initial.bullet_proposals)
    assert {item.source_bullet_id: item.proposed_text for item in initial.bullet_proposals} == profile.bullet_lookup()


def test_initial_report_detects_profile_title_mismatch(project_root, profile, analysis):
    initial = build_initial_resume_proposal(profile)
    report = build_resume_report(
        profile,
        analysis,
        initial,
        generated_filename="Vincent_Wenger_Initial_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        resume_title=initial_resume_title(profile),
    )
    title_check = next(
        check
        for check in _all_checks(report.searchability)
        if check.label == "The job title matches the resume profile title"
    )

    assert initial_resume_title(profile) == "Software Engineer"
    assert title_check.status == "fail"
    assert "Axiom Developer" in title_check.detail


def test_tailored_report_exact_target_title_passes(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        resume_title=analysis.target_title,
    )
    title_check = next(
        check
        for check in _all_checks(report.searchability)
        if check.label == "The job title matches the resume profile title"
    )

    assert title_check.status == "pass"



def test_skill_frequency_coverage_increases_job_match_without_changing_quality(
    project_root,
    profile,
    proposal,
):
    analysis = JobAnalysis(
        target_title="Software Engineer",
        target_company="Example Company",
        requirements=[
            JobRequirement(
                id="FREQ-1",
                category="technical_skill",
                priority="critical",
                requirement="Apply reporting frameworks",
                keywords=["reporting frameworks"],
            )
        ],
        ignored_boilerplate=[],
    )
    once = proposal.model_copy(deep=True)
    once.professional_summary = "Software engineer experienced with reporting frameworks."
    once.evidence_matches = [
        EvidenceMatch(
            requirement_id="FREQ-1",
            status="supported",
            evidence_ids=["NAS-01"],
            rationale="Supported for scoring test purposes.",
        )
    ]
    three_times = once.model_copy(deep=True)
    three_times.professional_summary = (
        "Software engineer experienced with reporting frameworks. "
        "Designed reporting frameworks and improved reporting frameworks."
    )
    job_description = (
        "The role designs reporting frameworks, governs reporting frameworks, "
        "and improves reporting frameworks."
    )

    initial = build_resume_report(
        profile,
        analysis,
        once,
        generated_filename="Candidate_Software_Engineer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        job_description=job_description,
        resume_title="Software Engineer",
    )
    updated = build_resume_report(
        profile,
        analysis,
        three_times,
        generated_filename="Candidate_Software_Engineer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        job_description=job_description,
        resume_title="Software Engineer",
    )

    initial_check = initial.hard_skills.subsections[0].checks[0]
    updated_check = updated.hard_skills.subsections[0].checks[0]
    assert initial_check.status == "warning"
    assert initial_check.score() == pytest.approx(33.333, rel=1e-3)
    assert updated_check.status == "pass"
    assert updated_check.score() == 100.0
    assert updated.hard_skills.score() > initial.hard_skills.score()
    assert updated.job_match_score() > initial.job_match_score()
    assert updated.overall_score() > initial.overall_score()
    assert updated.resume_quality_score() == pytest.approx(initial.resume_quality_score(), abs=0.2)


def test_report_exposes_decimal_weighted_scores(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        job_description=(
            "Axiom Axiom SQL SQL SQL data transformation testing collaboration."
        ),
    )

    assert isinstance(report.overall_score(), float)
    assert isinstance(report.job_match_score(), float)
    assert isinstance(report.resume_quality_score(), float)
    expected = round(
        report.searchability.score() * 0.15
        + report.hard_skills.score() * 0.25
        + report.soft_skills.score() * 0.08
        + report.content_quality.score() * 0.15
        + report.recruiter_tips.score() * 0.12
        + report.formatting.score() * 0.10
        + report.evidence_gaps.score() * 0.15,
        1,
    )
    assert report.overall_score() == expected


def test_evidence_gap_report_summarizes_support_and_locations(profile, analysis, proposal):
    analysis.requirements.extend(
        [
            JobRequirement(
                id="R4",
                category="technical_skill",
                priority="important",
                requirement="Administer Oracle RAC environments",
                keywords=["Oracle RAC"],
            ),
            JobRequirement(
                id="R5",
                category="technical_skill",
                priority="secondary",
                requirement="Develop COBOL batch applications",
                keywords=["COBOL"],
            ),
        ]
    )
    proposal.evidence_matches.extend(
        [
            EvidenceMatch(
                requirement_id="R4",
                status="partial",
                evidence_ids=["NAS-04"],
                rationale="The source supports Oracle work, but not Oracle RAC administration.",
            ),
            EvidenceMatch(
                requirement_id="R5",
                status="unsupported",
                evidence_ids=[],
                rationale="No verified COBOL evidence.",
            ),
        ]
    )
    proposal.unsupported_requirements = ["Develop COBOL batch applications"]
    proposal.candidate_questions = ["Have you administered Oracle RAC environments?"]

    summary, rows = build_evidence_gap_report(profile, analysis, proposal)
    by_id = {row.requirement_id: row for row in rows}

    assert summary.supported == 3
    assert summary.partial == 1
    assert summary.unsupported == 1
    assert summary.candidate_confirmations == 1
    assert by_id["R1"].appears_in_resume is True
    assert any("NAS-01" in location for location in by_id["R1"].evidence_locations)
    assert by_id["R4"].evidence_status == "partial"
    assert by_id["R4"].appears_in_resume is False
    assert "confirm" in by_id["R4"].recommended_action.casefold()
    assert by_id["R5"].evidence_status == "unsupported"
    assert "do not add" in by_id["R5"].recommended_action.casefold()


def test_evidence_matrix_is_inside_updated_report_workflow(project_root):
    template = (project_root / "templates" / "index.html").read_text(encoding="utf-8")
    source = (project_root / "resume_tailor" / "resume_report.py").read_text(encoding="utf-8")

    assert "section.name == 'Evidence & Gaps'" in template
    assert '"Evidence & Gaps": ("Review Job Alignment", "draft", "", "")' in source
    assert "Evidence was already reviewed in Step 3" in template

def test_evidence_gaps_is_scored_and_affects_overall_and_job_match(
    project_root,
    profile,
    analysis,
    proposal,
):
    supported = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )

    weakened_proposal = proposal.model_copy(deep=True)
    weakened_proposal.evidence_matches[0] = EvidenceMatch(
        requirement_id=weakened_proposal.evidence_matches[0].requirement_id,
        status="unsupported",
        evidence_ids=[],
        rationale="Evidence removed for scoring test.",
    )
    weakened = build_resume_report(
        profile,
        analysis,
        weakened_proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )

    assert supported.evidence_gaps.score() > weakened.evidence_gaps.score()
    assert supported.overall_score() > weakened.overall_score()
    assert supported.job_match_score() > weakened.job_match_score()
    assert supported.resume_quality_score() == weakened.resume_quality_score()


def test_evidence_gap_scores_supported_partial_no_and_unsupported_claims(
    project_root,
    profile,
    analysis,
    proposal,
):
    changed_analysis = analysis.model_copy(deep=True)
    changed_analysis.requirements.extend(
        [
            JobRequirement(
                id="R4",
                category="technical_skill",
                priority="important",
                requirement="Administer Oracle RAC",
                keywords=["Oracle RAC"],
            ),
            JobRequirement(
                id="R5",
                category="technical_skill",
                priority="secondary",
                requirement="Develop COBOL applications",
                keywords=["COBOL"],
            ),
            JobRequirement(
                id="R6",
                category="leadership",
                priority="secondary",
                requirement="Coach employees",
                keywords=["coach employees"],
            ),
        ]
    )
    changed_proposal = proposal.model_copy(deep=True)
    changed_proposal.evidence_matches.extend(
        [
            EvidenceMatch(
                requirement_id="R4",
                status="partial",
                evidence_ids=["NAS-04"],
                rationale="Oracle is supported, RAC administration is not fully supported.",
            ),
            EvidenceMatch(
                requirement_id="R5",
                status="unsupported",
                evidence_ids=[],
                rationale="No COBOL evidence.",
            ),
            EvidenceMatch(
                requirement_id="R6",
                status="unsupported",
                evidence_ids=[],
                rationale="The candidate did not confirm employee coaching.",
            ),
        ]
    )
    changed_proposal.skills.hard_skills.append("COBOL")
    answers = [
        CandidateAnswer(
            question_id="Q-R4",
            question="Have you administered Oracle RAC?",
            requirement_id="R4",
            answer_type="yes_no",
            yes_no=False,
        ),
        CandidateAnswer(
            question_id="Q-R6",
            question="Have you coached employees?",
            requirement_id="R6",
            answer_type="yes_no",
            yes_no=False,
        ),
    ]

    report = build_resume_report(
        profile,
        changed_analysis,
        changed_proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        candidate_answers=answers,
    )
    checks = {check.label: check for check in _all_checks(report.evidence_gaps)}

    assert checks["Administer Oracle RAC"].score() == 40.0
    assert checks["Develop COBOL applications"].score() == 0.0
    assert checks["Develop COBOL applications"].status == "fail"
    assert checks["Coach employees"].score() == 20.0
    assert "acknowledged gap" in checks["Coach employees"].detail.casefold()


def test_evidence_priority_weights_critical_requirements_more_heavily(
    project_root,
    profile,
    proposal,
):
    analysis = JobAnalysis(
        target_title="Software Engineer",
        requirements=[
            JobRequirement(
                id="CRIT",
                category="technical_skill",
                priority="critical",
                requirement="Use COBOL",
                keywords=["COBOL"],
            ),
            JobRequirement(
                id="SEC",
                category="leadership",
                priority="secondary",
                requirement="Communicate clearly",
                keywords=["communication"],
            ),
        ],
    )
    changed = proposal.model_copy(deep=True)
    changed.evidence_matches = [
        EvidenceMatch(
            requirement_id="CRIT",
            status="unsupported",
            evidence_ids=[],
            rationale="No evidence.",
        ),
        EvidenceMatch(
            requirement_id="SEC",
            status="supported",
            evidence_ids=["NAS-01"],
            rationale="Supported.",
        ),
    ]
    changed.skills.soft_skills.append("communication")

    report = build_resume_report(
        profile,
        analysis,
        changed,
        generated_filename="Candidate_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )

    # Critical unsupported and absent receives 10 (weight 3), while the
    # secondary supported requirement receives 100 (weight 1): 32.5 overall.
    assert report.evidence_gaps.score() == 32.5


def test_initial_resume_proposal_can_reuse_evidence_without_tailored_wording(profile, proposal):
    initial = build_initial_resume_proposal(profile, proposal)

    assert initial.professional_summary == profile.current_summary
    assert initial.evidence_matches == proposal.evidence_matches
    assert initial.candidate_questions == proposal.candidate_questions
    assert initial.bullet_proposals[0].proposed_text == profile.experiences[0].bullets[0].text


def test_confirmation_step_precedes_job_alignment_and_final_optimization(project_root):
    template = (project_root / "templates" / "index.html").read_text(encoding="utf-8")

    confirmation_position = template.index("<h2>Confirm Your Experience</h2>")
    generated_position = template.index("<h2>Review Job Alignment</h2>")
    final_position = template.index("<h2>Optimize &amp; Export</h2>")
    assert confirmation_position < generated_position < final_position
    assert "Create tailored resume" in template
    assert "Optimize &amp; Export" in template

def test_content_quality_contains_new_diagnostic_checks(project_root, profile, analysis, proposal):
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Vincent_Wenger_Axiom_Developer_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )

    assert [subsection.name for subsection in report.content_quality.subsections] == [
        "Data & Structure",
        "Semantic Match",
        "Grammar & Spelling",
        "Metric Integrity",
        "Readability",
        "Writing Style",
        "Content Focus",
    ]
    labels = {check.label for check in _all_checks(report.content_quality)}
    assert {
        "Core resume entities are complete",
        "The generated resume preserves the expected section hierarchy",
        "No common spelling, repeated-word, or punctuation errors were detected",
        "Every numeric claim is traceable to verified source evidence",
        "The resume has a recruiter-friendly readability level",
        "The resume avoids personal pronouns",
        "Past roles use consistent past-tense openings",
        "Supported priority requirements are represented without inventing experience",
    }.issubset(labels)


def test_semantic_match_scores_supported_but_unrepresented_requirement(
    project_root,
    profile,
    analysis,
    proposal,
):
    analysis.requirements.append(
        JobRequirement(
            id="SEM-1",
            category="responsibility",
            priority="important",
            requirement="Lead satellite telemetry calibration",
            keywords=["satellite telemetry calibration"],
        )
    )
    proposal.evidence_matches.append(
        EvidenceMatch(
            requirement_id="SEM-1",
            status="supported",
            evidence_ids=["NAS-01"],
            rationale="Supported for semantic-scoring test purposes.",
        )
    )

    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Candidate_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    check = next(
        item
        for item in report.content_quality.subsections[1].checks
        if item.label.startswith("SEM-1 semantic match")
    )

    assert check.status == "warning"
    assert check.score() == 70.0
    assert "not clearly represented" in check.detail


def test_language_and_style_checks_flag_common_errors(project_root, profile, analysis, proposal):
    changed = proposal.model_copy(deep=True)
    changed.professional_summary = "i analized teh teh reporting enviroment  ."

    report = build_resume_report(
        profile,
        analysis,
        changed,
        generated_filename="Candidate_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    by_label = {check.label: check for check in _all_checks(report.content_quality)}

    assert by_label["No common spelling, repeated-word, or punctuation errors were detected"].status == "fail"
    assert "analized → analyzed" in by_label["No common spelling, repeated-word, or punctuation errors were detected"].detail
    assert by_label["The resume avoids personal pronouns"].status == "warning"


def test_metric_integrity_flags_unverified_and_implausible_percentage(
    project_root,
    profile,
    analysis,
    proposal,
):
    changed = proposal.model_copy(deep=True)
    first = next(item for item in changed.bullet_proposals if item.include)
    first.proposed_text += " Increased throughput by 9001%."

    report = build_resume_report(
        profile,
        analysis,
        changed,
        generated_filename="Candidate_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
    )
    by_label = {check.label: check for check in _all_checks(report.content_quality)}

    assert by_label["Every numeric claim is traceable to verified source evidence"].status == "fail"
    assert "9001%" in by_label["Every numeric claim is traceable to verified source evidence"].detail
    assert by_label["Metrics are logically plausible"].status == "warning"


def test_configurable_one_page_limit_fails_for_two_page_resume(
    monkeypatch,
    project_root,
    profile,
    analysis,
    proposal,
):
    monkeypatch.setattr(
        "resume_tailor.resume_report._rendered_page_count",
        lambda document, *, exact=True: (2, "rendered"),
    )
    report = build_resume_report(
        profile,
        analysis,
        proposal,
        generated_filename="Candidate_Resume.docx",
        template_path=project_root / "data" / "resume_template_professional.docx",
        page_limit=1,
        exact_page_count=True,
    )
    check = next(
        item
        for item in _all_checks(report.formatting)
        if item.label == "The resume fits within the 1-page limit"
    )

    assert check.status == "fail"
    assert "exceeds" in check.detail
