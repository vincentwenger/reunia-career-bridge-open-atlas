from __future__ import annotations

from docx import Document

from products.resume_taylor.resume_tailor.docx_styles import (
    STYLE_BULLET,
    STYLE_CONTACT,
    STYLE_EDUCATION,
    STYLE_EMPLOYER_LINE,
    STYLE_NAME,
    STYLE_ROLE,
    STYLE_SECTION_HEADING,
    STYLE_SUMMARY,
    STYLE_TARGET_TITLE,
    clear_document_body,
    configure_resume_document,
)
from products.resume_taylor.resume_tailor.resume_pagination import (
    estimate_resume_pagination,
    rebalance_resume_pagination,
)
from products.resume_taylor.resume_tailor.resume_report import _formatting_sections


def _build_resume_with_bullet_count(bullet_count: int) -> Document:
    document = Document()
    configure_resume_document(
        document,
        career_stage="mid_career",
        visual_design="corporate",
    )
    clear_document_body(document)

    document.add_paragraph("Candidate Name", style=STYLE_NAME)
    document.add_paragraph("Data Engineer", style=STYLE_TARGET_TITLE)
    document.add_paragraph("Portland, OR", style=STYLE_CONTACT)
    document.add_paragraph("candidate@example.com | LinkedIn", style=STYLE_CONTACT)

    document.add_paragraph("Technical Profile", style=STYLE_SECTION_HEADING)
    document.add_paragraph(
        "Experienced engineer delivering financial data platforms, data pipelines, "
        "database upgrades, reporting controls, and production support.",
        style=STYLE_SUMMARY,
    )

    document.add_paragraph("Engineering Experience", style=STYLE_SECTION_HEADING)
    employer = document.add_paragraph(style=STYLE_EMPLOYER_LINE)
    employer.add_run("Example Financial Platform").bold = True
    employer.add_run("\t01/2015 - 05/2025")
    document.add_paragraph("Software Engineer", style=STYLE_ROLE)
    bullet_text = (
        "• Designed and implemented data-processing workflows across regulated systems "
        "while validating transformations and coordinating production releases."
    )
    for _ in range(bullet_count):
        document.add_paragraph(bullet_text, style=STYLE_BULLET)

    document.add_paragraph("Education and Certifications", style=STYLE_SECTION_HEADING)
    first = document.add_paragraph(style=STYLE_EDUCATION)
    first.add_run("Master of Science, Example University").bold = True
    first.add_run("\t09/2007")
    final = document.add_paragraph(style=STYLE_EDUCATION)
    final.add_run("Bachelor of Science, Example University").bold = True
    final.add_run("\t09/2003")
    return document


def _find_single_line_orphan_fixture() -> Document:
    for bullet_count in range(12, 40):
        document = _build_resume_with_bullet_count(bullet_count)
        estimate = estimate_resume_pagination(document)
        if estimate.has_orphan_final_page and estimate.last_page_substantive_lines == 1:
            return document
    raise AssertionError("Could not construct the expected single-line orphan fixture.")


def test_rebalance_compacts_single_line_final_page_without_shrinking_font_or_margins() -> None:
    document = _find_single_line_orphan_fixture()
    before = estimate_resume_pagination(document)
    original_font_size = document.styles[STYLE_BULLET].font.size.pt
    original_margins = [
        (section.top_margin.inches, section.bottom_margin.inches)
        for section in document.sections
    ]

    adjustment = rebalance_resume_pagination(document)
    after = estimate_resume_pagination(document)

    assert before.page_count == 2
    assert before.last_page_substantive_lines == 1
    assert adjustment.adjusted is True
    assert adjustment.mode in {"compact", "extra_compact"}
    assert after.page_count == 1
    assert after.has_orphan_final_page is False
    assert document.styles[STYLE_BULLET].font.size.pt == original_font_size
    assert [
        (section.top_margin.inches, section.bottom_margin.inches)
        for section in document.sections
    ] == original_margins


def test_rebalance_leaves_short_one_page_resume_unchanged() -> None:
    document = _build_resume_with_bullet_count(2)
    section_spacing_before = document.styles[STYLE_SECTION_HEADING].paragraph_format.space_before.pt

    adjustment = rebalance_resume_pagination(document)

    assert adjustment.mode == "none"
    assert adjustment.before.page_count == 1
    assert adjustment.after.page_count == 1
    assert document.styles[STYLE_SECTION_HEADING].paragraph_format.space_before.pt == section_spacing_before


def test_formatting_report_includes_orphan_final_page_quality_check() -> None:
    document = _find_single_line_orphan_fixture()

    subsections = _formatting_sections(
        document,
        inspection_note=None,
        page_limit=2,
        exact_page_count=False,
    )
    page_setup = next(section for section in subsections if section.name == "Page Setup")
    check = next(
        item for item in page_setup.checks if item.label == "The final page is not nearly empty"
    )

    assert check.status == "fail"
    assert "1 substantive line" in check.detail
