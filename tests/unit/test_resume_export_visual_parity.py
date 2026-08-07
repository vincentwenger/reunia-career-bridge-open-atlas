from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether, Paragraph

from products.resume_taylor.resume_tailor.docx_export import export_resume_docx
from products.resume_taylor.resume_tailor.docx_styles import (
    STYLE_BULLET,
    STYLE_NAME,
    STYLE_SECTION_HEADING,
    STYLE_SUMMARY,
    STYLE_TARGET_TITLE,
    configure_resume_document,
)
from products.resume_taylor.resume_tailor.models import (
    ApprovedResume,
    CandidateProfile,
    ContactInfo,
    EducationItem,
    Experience,
    ResumeBullet,
    SkillSet,
    VerifiedSkills,
)
from products.resume_taylor.resume_tailor.pdf_export import (
    _add_education,
    _add_experience,
    _build_styles,
)


TEMPLATE = (
    Path(__file__).resolve().parents[2]
    / "products"
    / "resume_taylor"
    / "data"
    / "resume_template_professional.docx"
)


def _profile() -> CandidateProfile:
    return CandidateProfile(
        name="Candidate Name",
        contact=ContactInfo(
            location="Portland, OR",
            phone="555-0100",
            email="candidate@example.com",
        ),
        current_summary="Experienced data engineer.",
        skills=VerifiedSkills(),
        education=[
            EducationItem(
                credential="Master of Science (M.S.), Engineering",
                institution="Example University",
                location="Paris, France",
                date="09/2007",
                detail="Emphasis in Information Systems Engineering",
            )
        ],
        experiences=[
            Experience(
                id="EXP-001",
                employer="Harbor Mutual",
                location="Paris, France",
                dates="08/2010 - 08/2011",
                title="IT Auditor",
                bullets=[ResumeBullet(id="EXP-001-B01", text="Audited IT controls.")],
            ),
            Experience(
                id="EXP-002",
                employer="Capgemini",
                location="Paris, France",
                dates="10/2007 - 08/2010",
                title="QA Engineer",
                bullets=[ResumeBullet(id="EXP-002-B01", text="Performed quality assurance testing.")],
            ),
        ],
    )


def _approved() -> ApprovedResume:
    return ApprovedResume(
        target_title="Data Engineer",
        professional_summary="Experienced data engineer.",
        skills=SkillSet(),
        bullets_by_experience={
            "EXP-001": ["Audited IT controls."],
            "EXP-002": ["Performed quality assurance testing."],
        },
    )


def test_mid_career_docx_uses_bold_employers_normal_tracking_and_matching_spacing() -> None:
    payload = export_resume_docx(
        TEMPLATE,
        _profile(),
        _approved(),
        career_stage="mid_career",
        resume_format="technical",
        visual_design="corporate",
    )
    document = Document(BytesIO(payload))

    employers = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(("Harbor Mutual", "Capgemini"))
    ]
    assert len(employers) == 2
    for paragraph in employers:
        employer_run = paragraph.runs[0]
        assert employer_run.bold is True
        assert employer_run.underline is not True

    assert employers[0].paragraph_format.space_before.pt == 0
    assert employers[1].paragraph_format.space_before.pt == 6

    for style_name in (STYLE_NAME, STYLE_TARGET_TITLE, STYLE_SECTION_HEADING):
        r_pr = document.styles[style_name]._element.get_or_add_rPr()
        assert r_pr.find(qn("w:spacing")) is None

    bullet_format = document.styles[STYLE_BULLET].paragraph_format
    assert abs(bullet_format.left_indent.inches - 0.22) < 0.001
    assert abs(bullet_format.first_line_indent.inches + 0.17) < 0.001
    assert abs(document.styles[STYLE_SUMMARY].paragraph_format.line_spacing.pt - 11.8) < 0.01

    education_detail = next(
        paragraph
        for paragraph in document.paragraphs
        if "Emphasis in Information Systems Engineering" in paragraph.text
    )
    assert education_detail.text.startswith("• ")


def test_mid_career_pdf_matches_bold_employer_bullet_indentation_and_education_marker() -> None:
    document = Document()
    theme = configure_resume_document(
        document,
        career_stage="mid_career",
        visual_design="corporate",
    )
    styles = _build_styles(theme)
    story: list = []

    _add_experience(
        story,
        _profile(),
        _approved(),
        theme,
        styles,
        heading="Engineering Experience",
        usable_width=500,
    )
    headings = [item for item in story if isinstance(item, KeepTogether)]
    first_employer = headings[0]._content[0]._cellvalues[0][0]
    assert isinstance(first_employer, Paragraph)
    assert first_employer.getPlainText().startswith("Harbor Mutual, Paris, France")
    assert "Bold" in first_employer.frags[0].fontName

    assert abs(styles["bullet"].leftIndent - (0.22 * inch)) < 0.01
    assert abs(styles["bullet"].firstLineIndent - (-0.17 * inch)) < 0.01
    assert abs(styles["body"].leading - 11.8) < 0.01

    education_story: list = []
    _add_education(
        education_story,
        _profile(),
        theme,
        styles,
        heading="Education and Certifications",
        usable_width=500,
    )
    education_tables = [
        item._content[0]
        for item in education_story
        if isinstance(item, KeepTogether)
    ]
    assert education_tables
    assert abs(education_tables[0]._colWidths[0] - (500 * 0.88)) < 0.01

    details = [
        item
        for item in education_story
        if isinstance(item, Paragraph)
        and "Emphasis in Information Systems Engineering" in item.getPlainText()
    ]
    assert len(details) == 1
    assert details[0].getPlainText().startswith("• ")
