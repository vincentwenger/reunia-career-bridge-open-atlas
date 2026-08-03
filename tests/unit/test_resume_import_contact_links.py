from __future__ import annotations

from io import BytesIO
import unittest

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from products.resume_taylor.resume_tailor.models import (
    CandidateProfile,
    ContactInfo,
    VerifiedSkills,
)
from products.resume_taylor.resume_tailor.resume_import import (
    extract_resume_text,
    inherit_professional_contact_urls,
    professional_contact_urls,
    restore_professional_contact_urls,
)


def _add_hyperlink(paragraph, label: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        RT.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _empty_profile() -> CandidateProfile:
    return CandidateProfile(
        name="Vincent Wenger",
        contact=ContactInfo(location="Portland, OR", phone="", email="vincent@example.com"),
        current_summary="Senior software engineer with banking technology experience.",
        skills=VerifiedSkills(),
        education=[],
        experiences=[],
    )


class ResumeImportContactLinkTests(unittest.TestCase):
    def test_word_header_hyperlinks_are_included_in_extracted_text(self) -> None:
        document = Document()
        header = document.sections[0].header.paragraphs[0]
        header.add_run("Vincent Wenger | ")
        _add_hyperlink(
            header,
            "LinkedIn",
            "https://www.linkedin.com/in/vincentwenger/",
        )
        header.add_run(" | ")
        _add_hyperlink(
            header,
            "GitHub",
            "https://github.com/vincentwenger",
        )
        document.add_paragraph(
            "Senior Software Engineer with fifteen years of experience in banking technology."
        )

        buffer = BytesIO()
        document.save(buffer)
        extracted = extract_resume_text(buffer.getvalue(), "resume.docx")

        self.assertIn("Vincent Wenger", extracted)
        self.assertIn("https://www.linkedin.com/in/vincentwenger/", extracted)
        self.assertIn("https://github.com/vincentwenger", extracted)

    def test_professional_urls_are_classified_from_resume_text(self) -> None:
        urls = professional_contact_urls(
            "LinkedIn: linkedin.com/in/vincentwenger GitHub: www.github.com/vincentwenger"
        )

        self.assertEqual(
            urls["linkedin_url"],
            "https://linkedin.com/in/vincentwenger",
        )
        self.assertEqual(
            urls["github_url"],
            "https://www.github.com/vincentwenger",
        )

    def test_missing_structured_urls_are_restored_without_overwriting_existing_values(self) -> None:
        profile = _empty_profile()
        profile.contact.linkedin_url = "https://linkedin.example/existing"

        restored = restore_professional_contact_urls(
            profile,
            "LinkedIn https://linkedin.com/in/replacement GitHub https://github.com/vincentwenger",
        )

        self.assertEqual(
            restored.contact.linkedin_url,
            "https://linkedin.example/existing",
        )
        self.assertEqual(
            restored.contact.github_url,
            "https://github.com/vincentwenger",
        )
        self.assertEqual(profile.contact.github_url, "")

    def test_profile_copy_inherits_only_missing_professional_links(self) -> None:
        source = _empty_profile()
        source.contact.linkedin_url = "https://linkedin.com/in/source"
        source.contact.github_url = "https://github.com/source"

        target = _empty_profile()
        target.contact.linkedin_url = "https://linkedin.example/application-specific"
        restored = inherit_professional_contact_urls(target, source)

        self.assertEqual(
            restored.contact.linkedin_url,
            "https://linkedin.example/application-specific",
        )
        self.assertEqual(restored.contact.github_url, "https://github.com/source")
        self.assertEqual(target.contact.github_url, "")

    def test_application_baseline_template_displays_clickable_professional_links(self) -> None:
        with open(
            "products/resume_taylor/templates/application_builder/index.html",
            encoding="utf-8",
        ) as template_file:
            template = template_file.read()

        self.assertIn("profile.contact.linkedin_url", template)
        self.assertIn("profile.contact.github_url", template)
        self.assertIn('class="resume-paper-contact-links"', template)
        self.assertIn('rel="noopener noreferrer"', template)

    def test_baseline_resume_template_displays_clickable_professional_links(self) -> None:
        with open(
            "products/resume_taylor/templates/application_builder/career_translation.html",
            encoding="utf-8",
        ) as template_file:
            template = template_file.read()

        self.assertIn("source_profile.contact.linkedin_url", template)
        self.assertIn("source_profile.contact.github_url", template)
        self.assertIn('class="career-translation-contact-links"', template)
        self.assertIn('rel="noopener noreferrer"', template)


if __name__ == "__main__":
    unittest.main()
